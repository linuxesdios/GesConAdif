#!/usr/bin/env python3
"""
Controlador de Facturas Directas - Sistema ADIF
Gestiona facturas directas con interfaz estilo verde corporativo ADIF
"""

import sys
import os
import json
import shutil
from datetime import datetime
from typing import Dict, List, Any, Optional
import logging

logger = logging.getLogger(__name__)
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

# Importaciones para generar documentos Word
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx no está instalado. Instalar con: pip install python-docx")

class ControladorFacturasDirectas:
    """Controlador principal para gestión de facturas directas"""
    
    def __init__(self, parent=None):
        self.parent = parent
        # Usar el sistema de rutas centralizado
        from .controlador_routes import rutas
        self.directorio_base = rutas.get_base_path()
        self.archivo_json = rutas.get_ruta_facturas_directas()
        self.carpeta_pdfs = os.path.join(self.directorio_base, "pdfactura directa")
        self._inicializar_archivos()
    
    def _obtener_directorio_base(self):
        """Obtener directorio base donde está BaseDatos.json - Compatible con exe"""
        try:
            # Para PyInstaller: usar sys._MEIPASS si existe
            if hasattr(sys, '_MEIPASS'):
                # En exe: buscar en el directorio donde está el exe
                exe_dir = os.path.dirname(sys.executable)
                # Verificar si BaseDatos.json existe en el directorio del exe
                if os.path.exists(os.path.join(exe_dir, "BaseDatos.json")):
                    return exe_dir
                # Si no existe, usar el directorio actual
                return os.getcwd()
            else:
                # En desarrollo: usar la estructura normal
                if self.parent:
                    current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                else:
                    current_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                return current_dir
        except Exception as e:
            logger.warning(f"Error obteniendo directorio base: {e}")
            return os.getcwd()
    
    def _inicializar_archivos(self):
        """Inicializar archivo JSON y carpeta de PDFs si no existen"""
        try:
            # Crear archivo JSON si no existe
            if not os.path.exists(self.archivo_json):
                datos_iniciales = {
                    "version": "1.1",
                    "fecha_creacion": datetime.now().isoformat(),
                    "facturas": [],
                    "configuracion": {
                        "categorias": ["Agua", "Vegetal", "Limpieza", "Actuaciones mantenimiento", "Otras"],
                        "estados": ["Emitida", "Tramitada", "Pagada", "Con deficiencias"],
                        "ultimo_id": 0
                    }
                }
                with open(self.archivo_json, 'w', encoding='utf-8') as f:
                    json.dump(datos_iniciales, f, ensure_ascii=False, indent=4)
                logger.info(f"Archivo JSON creado: {self.archivo_json}")
            
            # Crear carpeta de PDFs si no existe
            if not os.path.exists(self.carpeta_pdfs):
                os.makedirs(self.carpeta_pdfs, exist_ok=True)
                logger.info(f"Carpeta PDFs creada: {self.carpeta_pdfs}")
                
        except Exception as e:
            logger.error(f"Error inicializando archivos: {e}")
    
    def mostrar_popup_principal(self):
        """Mostrar popup principal de Facturas Directas"""
        try:
            dialogo = DialogoFacturasDirectas(self.parent, self)
            return dialogo.exec_()
        except Exception as e:
            QMessageBox.critical(self.parent, "Error", f"Error abriendo Facturas Directas: {str(e)}")
            return False
    
    def leer_datos_json(self) -> Dict:
        """Leer datos del archivo JSON"""
        try:
            with open(self.archivo_json, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Error leyendo JSON: {e}")
            return {"facturas": [], "configuracion": {"categorias": [], "ultimo_id": 0}}
    
    def guardar_datos_json(self, datos: Dict) -> bool:
        """Guardar datos en el archivo JSON"""
        try:
            with open(self.archivo_json, 'w', encoding='utf-8') as f:
                json.dump(datos, f, ensure_ascii=False, indent=4)
            return True
        except Exception as e:
            logger.error(f"Error guardando JSON: {e}")
            return False
    
    def agregar_factura(self, datos_factura: Dict) -> bool:
        """Agregar nueva factura"""
        try:
            datos = self.leer_datos_json()
            
            # Generar ID único
            datos["configuracion"]["ultimo_id"] += 1
            nuevo_id = datos["configuracion"]["ultimo_id"]
            
            # Crear factura con timestamp y todos los campos
            nueva_factura = {
                "id": nuevo_id,
                "fecha_creacion": datetime.now().isoformat(),
                "importe": datos_factura.get("importe", 0.0),
                "categoria": datos_factura.get("categoria", ""),
                "localidad": datos_factura.get("localidad", ""),
                "empresa": datos_factura.get("empresa", ""),
                "cif": datos_factura.get("cif", ""),
                "identificador_especial": datos_factura.get("identificador_especial", ""),
                "identificacion_admycont": datos_factura.get("identificacion_admycont", ""),
                "fecha_validacion": datos_factura.get("fecha_validacion", datetime.now().date().isoformat()),
                "estado": datos_factura.get("estado", "Emitida"),
                "comentarios": datos_factura.get("comentarios", ""),
                "gped": datos_factura.get("gped", ""),
                "archivo_pdf": datos_factura.get("archivo_pdf", ""),
                "activa": True
            }
            
            datos["facturas"].append(nueva_factura)
            
            if self.guardar_datos_json(datos):
                logger.info(f"Factura agregada con ID: {nuevo_id}")
                return True
            return False
            
        except Exception as e:
            logger.error(f"Error agregando factura: {e}")
            return False
    
    def obtener_facturas(self) -> List[Dict]:
        """Obtener lista de facturas activas"""
        try:
            datos = self.leer_datos_json()
            return [f for f in datos.get("facturas", []) if f.get("activa", True)]
        except Exception as e:
            logger.error(f"Error obteniendo facturas: {e}")
            return []
    
    def eliminar_factura(self, factura_id: int) -> bool:
        """Eliminar factura (marcar como inactiva)"""
        try:
            datos = self.leer_datos_json()
            
            for factura in datos.get("facturas", []):
                if factura.get("id") == factura_id:
                    factura["activa"] = False
                    factura["fecha_eliminacion"] = datetime.now().isoformat()
                    
                    if self.guardar_datos_json(datos):
                        logger.info(f"Factura {factura_id} eliminada")
                        return True
                    break
            return False
            
        except Exception as e:
            logger.error(f"Error eliminando factura: {e}")
            return False
    
    def actualizar_estado_factura(self, factura_id: int, nuevo_estado: str) -> bool:
        """Actualizar estado de una factura específica"""
        try:
            datos = self.leer_datos_json()
            
            for factura in datos.get("facturas", []):
                if factura.get("id") == factura_id and factura.get("activa", True):
                    factura["estado"] = nuevo_estado
                    factura["fecha_modificacion_estado"] = datetime.now().isoformat()
                    
                    if self.guardar_datos_json(datos):
                        logger.info(f"Estado de factura {factura_id} actualizado a: {nuevo_estado}")
                        return True
                    break
            return False
            
        except Exception as e:
            logger.error(f"Error actualizando estado: {e}")
            return False
    
    def generar_nombre_pdf(self, id_factura: int, empresa: str, importe: float) -> str:
        """Generar nombre de PDF con formato: id_empresa_importe.pdf"""
        try:
            # Limpiar nombre de empresa (quitar caracteres especiales)
            empresa_limpia = "".join(c for c in empresa if c.isalnum() or c in (' ', '-', '_')).strip()
            empresa_limpia = empresa_limpia.replace(' ', '_')
            if not empresa_limpia:
                empresa_limpia = "SinEmpresa"
            
            # Formatear importe (reemplazar coma por punto para evitar problemas)
            importe_str = f"{importe:.2f}".replace('.', ',')
            
            # Generar nombre final
            nombre_pdf = f"{id_factura}_{empresa_limpia}_{importe_str}.pdf"
            
            return nombre_pdf
            
        except Exception as e:
            logger.error(f"Error generando nombre PDF: {e}")
            return f"{id_factura}_factura.pdf"
    
    def gestionar_pdf(self, archivo_origen: str, datos_factura: Dict) -> str:
        """Copiar PDF a la carpeta de facturas directas con nombre formato id_empresa_importe.pdf"""
        try:
            if not os.path.exists(archivo_origen):
                raise FileNotFoundError(f"Archivo origen no existe: {archivo_origen}")
            
            # Generar nombre con formato específico
            id_factura = datos_factura.get("id", 0)
            empresa = datos_factura.get("empresa", "SinEmpresa")
            importe = datos_factura.get("importe", 0.0)
            
            nombre_pdf = self.generar_nombre_pdf(id_factura, empresa, importe)
            archivo_destino = os.path.join(self.carpeta_pdfs, nombre_pdf)
            
            # Si el archivo ya existe, eliminarlo (reemplazar)
            if os.path.exists(archivo_destino):
                os.remove(archivo_destino)
            
            # Copiar archivo
            shutil.copy2(archivo_origen, archivo_destino)
            logger.info(f"PDF copiado con formato: {nombre_pdf}")
            
            return nombre_pdf
            
        except Exception as e:
            logger.error(f"Error gestionando PDF: {e}")
            return ""
    
    def generar_informe_word(self, facturas: List[Dict]) -> str:
        """Generar informe de facturas en formato Word (.docx)"""
        try:
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx no está disponible")
            
            # Crear documento
            doc = Document()
            
            # === PÁGINA INICIAL ===
            # Título principal
            titulo = doc.add_heading('', 0)
            titulo_run = titulo.runs[0] if titulo.runs else titulo.add_run()
            titulo.text = "INFORME DE FACTURAS DIRECTAS"
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titulo_run.font.size = Pt(20)
            titulo_run.font.color.rgb = RGBColor(46, 125, 50)  # Verde ADIF
            titulo_run.bold = True
            
            # Subtítulo
            subtitulo = doc.add_heading('Sistema de Administración ADIF', 2)
            subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitulo.runs[0].font.color.rgb = RGBColor(27, 94, 32)  # Verde oscuro ADIF
            
            # Espaciado
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Información general
            fecha_actual = datetime.now().strftime("%d/%m/%Y %H:%M")
            total_facturas = len(facturas)
            total_importe = sum(f.get('importe', 0) for f in facturas)
            
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(f"""
📅 Fecha de generación: {fecha_actual}
📊 Total de facturas: {total_facturas}
💰 Importe total: {total_importe:.2f} €
            """.strip())
            info_run.font.size = Pt(12)
            info_run.font.color.rgb = RGBColor(56, 142, 60)  # Verde info ADIF
            
            # Resumen por estado
            estados_resumen = {}
            for factura in facturas:
                estado = factura.get('estado', 'Sin estado')
                estados_resumen[estado] = estados_resumen.get(estado, 0) + 1
            
            doc.add_paragraph()
            resumen_para = doc.add_heading('📈 Resumen por Estado', 3)
            resumen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for estado, cantidad in estados_resumen.items():
                estado_para = doc.add_paragraph(f"• {estado}: {cantidad} facturas")
                estado_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                estado_para.runs[0].font.size = Pt(11)
            
            # Salto de página
            doc.add_page_break()
            
            # === PÁGINA DE TABLA ===
            # Título de la tabla
            tabla_titulo = doc.add_heading('📋 Detalle de Facturas', 1)
            tabla_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tabla_titulo.runs[0].font.color.rgb = RGBColor(46, 125, 50)
            
            if not facturas:
                doc.add_paragraph("No hay facturas para mostrar.")
                return self._guardar_documento(doc)
            
            # Crear tabla
            columnas = ['ID', 'Empresa', 'CIF', 'Importe', 'Estado', 'Categoría', 'Localidad', 'Fecha Valid.']
            tabla = doc.add_table(rows=1, cols=len(columnas))
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabla.style = 'Light Grid Accent 1'
            
            # Configurar encabezados
            header_cells = tabla.rows[0].cells
            for i, columna in enumerate(columnas):
                header_cells[i].text = columna
                # Formatear encabezados
                para = header_cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                # Color de fondo verde ADIF para encabezados
                shading_elm = self._crear_sombreado(header_cells[i], RGBColor(46, 125, 50))
            
            # Añadir filas de datos
            for factura in facturas:
                row_cells = tabla.add_row().cells
                
                # Preparar datos
                datos_fila = [
                    str(factura.get('id', '')),
                    factura.get('empresa', ''),
                    factura.get('cif', ''),
                    f"{factura.get('importe', 0):.2f} €",
                    factura.get('estado', ''),
                    factura.get('categoria', ''),
                    factura.get('localidad', ''),
                    factura.get('fecha_validacion', '')
                ]
                
                # Llenar celdas
                for i, dato in enumerate(datos_fila):
                    row_cells[i].text = dato
                    # Formatear celdas
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    
                    # Colorear según estado (columna 4)
                    if i == 4:  # Columna Estado
                        estado = factura.get('estado', '')
                        if estado == 'Pagada':
                            run.font.color.rgb = RGBColor(76, 175, 80)  # Verde
                        elif estado == 'Con deficiencias':
                            run.font.color.rgb = RGBColor(244, 67, 54)  # Rojo
                        elif estado == 'Tramitada':
                            run.font.color.rgb = RGBColor(255, 152, 0)  # Naranja
            
            # Ajustar ancho de columnas
            for i, column in enumerate(tabla.columns):
                if i == 0:  # ID
                    column.width = Inches(0.5)
                elif i == 1:  # Empresa
                    column.width = Inches(1.5)
                elif i == 2:  # CIF
                    column.width = Inches(1.0)
                elif i == 3:  # Importe
                    column.width = Inches(0.8)
                else:
                    column.width = Inches(1.0)
            
            # Footer con información adicional
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(f"\nInforme generado por Sistema ADIF - {fecha_actual}")
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(117, 117, 117)
            footer_run.italic = True
            
            return self._guardar_documento(doc)
            
        except Exception as e:
            logger.error(f"Error generando informe Word: {e}")
            return ""
    
    def _crear_sombreado(self, cell, color: RGBColor):
        """Crear sombreado para celda de tabla"""
        try:
            shading_elm = OxmlElement(qn('w:shd'))
            shading_elm.set(qn('w:fill'), f"{color.r:02x}{color.g:02x}{color.b:02x}")
            cell._tc.get_or_add_tcPr().append(shading_elm)
            return shading_elm
        except:
            return None
    
    def _guardar_documento(self, doc) -> str:
        """Guardar documento y devolver ruta"""
        try:
            # Generar nombre de archivo único
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nombre_archivo = f"Informe_Facturas_ADIF_{timestamp}.docx"
            ruta_archivo = os.path.join(self.directorio_base, nombre_archivo)
            
            # Guardar documento
            doc.save(ruta_archivo)
            logger.info(f"Informe guardado: {ruta_archivo}")
            return ruta_archivo
            
        except Exception as e:
            logger.error(f"Error guardando documento: {e}")
            return ""
    
    def eliminar_pdf(self, nombre_archivo: str) -> bool:
        """Eliminar PDF de la carpeta de facturas directas"""
        try:
            archivo_path = os.path.join(self.carpeta_pdfs, nombre_archivo)
            if os.path.exists(archivo_path):
                os.remove(archivo_path)
                logger.info(f"PDF eliminado: {nombre_archivo}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error eliminando PDF: {e}")
            return False


class DialogoFacturasDirectas(QDialog):
    """Popup principal para gestión de facturas directas con estilo ADIF Verde Corporativo"""
    
    def __init__(self, parent=None, controlador=None):
        super().__init__(parent)
        self.controlador = controlador
        self.setWindowTitle("Facturas Directas - ADIF")
        self.setFixedSize(650, 450)
        self.setWindowFlags(Qt.Dialog | Qt.WindowCloseButtonHint)
        self.setup_ui()
        self.apply_adif_style()
    
    def setup_ui(self):
        """Configurar la interfaz simplificada del popup"""
        layout = QVBoxLayout(self)
        layout.setSpacing(10)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Título simple
        title_label = QLabel("💰 Facturas Directas")
        title_label.setObjectName("title_label")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Información de estado
        if self.controlador:
            try:
                facturas = self.controlador.obtener_facturas()
                total = sum(f.get('importe', 0) for f in facturas)
                info_text = f"{len(facturas)} facturas • Total: {total:.2f} €"
            except:
                info_text = "Sistema listo"
        else:
            info_text = "Sistema listo"
            
        info_label = QLabel(info_text)
        info_label.setObjectName("info_label")
        info_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(info_label)
        
        layout.addSpacing(10)
        
        # Botones principales simples
        buttons_data = [
            ("📝 Crear factura", self.crear_factura),
            ("✏️ Editar factura", self.editar_factura),
            ("🗑️ Borrar factura", self.borrar_factura),
            ("📊 Resumen facturación", self.resumen_facturacion),
            ("📄 Informe facturación", self.informe_facturacion)
        ]
        
        for text, callback in buttons_data:
            btn = QPushButton(text)
            btn.setObjectName("main_button")
            btn.setFixedHeight(35)
            btn.clicked.connect(callback)
            layout.addWidget(btn)
        
        layout.addSpacing(15)
        
        # Botón Cerrar
        btn_cerrar = QPushButton("Cerrar")
        btn_cerrar.setObjectName("close_button")
        btn_cerrar.setFixedHeight(30)
        btn_cerrar.clicked.connect(self.close)
        layout.addWidget(btn_cerrar)
    
    def apply_adif_style(self):
        """Aplicar estilo ADIF Verde Corporativo simplificado"""
        style = """
        QDialog {
            background-color: #f8f9fa;
            border: 2px solid #2E7D32;
            border-radius: 10px;
            font-family: 'Segoe UI', Arial, sans-serif;
        }
        
        QLabel#title_label {
            font-size: 12pt;
            font-weight: bold;
            color: white;
            background-color: #2E7D32;
            padding: 10px;
            border-radius: 6px;
            margin: 5px;
        }
        
        QLabel#info_label {
            font-size: 8pt;
            color: #2E7D32;
            background-color: rgba(76, 175, 80, 0.1);
            padding: 6px;
            border: 1px solid #4CAF50;
            border-radius: 4px;
        }
        
        QPushButton#main_button {
            font-size: 9pt;
            font-weight: 500;
            color: white;
            background-color: #4CAF50;
            border: 1px solid #2E7D32;
            border-radius: 5px;
            padding: 8px;
        }
        
        QPushButton#main_button:hover {
            background-color: #66BB6A;
            border: 2px solid #2E7D32;
        }
        
        QPushButton#main_button:pressed {
            background-color: #2E7D32;
        }
        
        QPushButton#close_button {
            font-size: 8pt;
            color: #2E7D32;
            background-color: #ffffff;
            border: 2px solid #2E7D32;
            border-radius: 5px;
            padding: 6px;
        }
        
        QPushButton#close_button:hover {
            background-color: #e8f5e8;
        }
        """
        self.setStyleSheet(style)
    
    def crear_factura(self):
        """Mostrar popup para crear nueva factura"""
        dialogo = DialogoCrearFactura(self, self.controlador)
        if dialogo.exec_() == QDialog.Accepted:
            QMessageBox.information(self, "Factura creada", "Factura creada exitosamente")
    
    def editar_factura(self):
        """Editar factura existente"""
        facturas = self.controlador.obtener_facturas() if self.controlador else []
        if not facturas:
            QMessageBox.information(self, "Sin facturas", "No hay facturas disponibles para editar")
            return
        
        # Mostrar lista de facturas para seleccionar (con nuevos campos)
        items = [f"ID: {f['id']} - {f.get('empresa', 'Sin empresa')} - {f['categoria']} - {f.get('estado', 'N/A')} - {f['importe']}€" 
                for f in facturas]
        
        item, ok = QInputDialog.getItem(self, "Editar factura", 
                                      "Seleccione factura a editar:", items, 0, False)
        if ok and item:
            factura_id = int(item.split(" - ")[0].replace("ID: ", ""))
            factura_data = next(f for f in facturas if f['id'] == factura_id)
            
            dialogo = DialogoCrearFactura(self, self.controlador, factura_data)
            if dialogo.exec_() == QDialog.Accepted:
                QMessageBox.information(self, "Factura editada", "Factura editada exitosamente")
    
    def borrar_factura(self):
        """Borrar factura con confirmación"""
        facturas = self.controlador.obtener_facturas() if self.controlador else []
        if not facturas:
            QMessageBox.information(self, "Sin facturas", "No hay facturas disponibles para borrar")
            return
        
        # Mostrar lista de facturas para seleccionar (con nuevos campos)
        items = [f"ID: {f['id']} - {f.get('empresa', 'Sin empresa')} - {f['categoria']} - {f.get('estado', 'N/A')} - {f['importe']}€" 
                for f in facturas]
        
        item, ok = QInputDialog.getItem(self, "Borrar factura", 
                                      "Seleccione factura a borrar:", items, 0, False)
        if ok and item:
            reply = QMessageBox.question(self, "Confirmar borrado", 
                                       f"¿Está seguro de que desea borrar la factura:\n{item}?",
                                       QMessageBox.Yes | QMessageBox.No,
                                       QMessageBox.No)
            if reply == QMessageBox.Yes:
                factura_id = int(item.split(" - ")[0].replace("ID: ", ""))
                if self.controlador and self.controlador.eliminar_factura(factura_id):
                    QMessageBox.information(self, "Factura borrada", "Factura borrada exitosamente")
                else:
                    QMessageBox.critical(self, "Error", "No se pudo borrar la factura")
    
    def resumen_facturacion(self):
        """Mostrar resumen de facturación en tabla"""
        facturas = self.controlador.obtener_facturas() if self.controlador else []
        if not facturas:
            QMessageBox.information(self, "Sin facturas", "No hay facturas disponibles")
            return
        
        # Crear diálogo con tabla
        dialogo_tabla = DialogoTablaResumen(self, facturas)
        dialogo_tabla.exec_()
    
    def informe_facturacion(self):
        """Generar informe de facturación en Word"""
        try:
            if not self.controlador:
                QMessageBox.critical(self, "Error", "Controlador no disponible")
                return
                
            # Obtener facturas
            facturas = self.controlador.obtener_facturas()
            if not facturas:
                QMessageBox.information(self, "Sin facturas", "No hay facturas para generar el informe")
                return
            
            # Verificar que python-docx esté disponible
            if not DOCX_AVAILABLE:
                QMessageBox.critical(self, "Dependencia faltante", 
                                   "Para generar informes Word necesitas instalar python-docx:\n\n"
                                   "pip install python-docx\n\n"
                                   "Después reinicia la aplicación.")
                return
            
            # Mostrar diálogo de progreso
            progress = QProgressDialog("Generando informe Word...", "Cancelar", 0, 100, self)
            progress.setWindowModality(Qt.WindowModal)
            progress.setWindowTitle("Generando Informe")
            progress.setValue(10)
            
            # Generar informe
            progress.setValue(50)
            ruta_archivo = self.controlador.generar_informe_word(facturas)
            progress.setValue(90)
            
            if ruta_archivo and os.path.exists(ruta_archivo):
                progress.setValue(100)
                progress.close()
                
                # Preguntar si abrir el archivo
                reply = QMessageBox.question(self, "Informe generado", 
                                           f"📄 Informe generado exitosamente:\n\n{os.path.basename(ruta_archivo)}\n\n"
                                           f"📁 Ubicación: {os.path.dirname(ruta_archivo)}\n\n"
                                           "¿Desea abrir el archivo?",
                                           QMessageBox.Yes | QMessageBox.No,
                                           QMessageBox.Yes)
                
                if reply == QMessageBox.Yes:
                    try:
                        # Abrir archivo con aplicación predeterminada
                        if sys.platform.startswith('win'):
                            os.startfile(ruta_archivo)
                        elif sys.platform.startswith('darwin'):  # macOS
                            os.system(f'open "{ruta_archivo}"')
                        else:  # Linux
                            os.system(f'xdg-open "{ruta_archivo}"')
                    except Exception as e:
                        QMessageBox.information(self, "Archivo generado", 
                                              f"Informe guardado en:\n{ruta_archivo}")
            else:
                progress.close()
                QMessageBox.critical(self, "Error", "No se pudo generar el informe Word")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error generando informe: {str(e)}")


class DialogoCrearFactura(QDialog):
    """Dialog para crear/editar facturas con campos específicos"""
    
    def __init__(self, parent=None, controlador=None, datos_factura=None):
        super().__init__(parent)
        self.controlador = controlador
        self.datos_factura = datos_factura  # Para edición
        self.archivo_pdf_seleccionado = ""
        
        self.setWindowTitle("Crear Factura" if not datos_factura else "Editar Factura")
        self.setFixedSize(650, 700)  # Ampliado para nuevos campos
        self.setWindowFlags(Qt.Dialog | Qt.WindowCloseButtonHint)
        self.setup_ui()
        self.apply_adif_style()
        
        if datos_factura:
            self.cargar_datos_existentes()
    
    def setup_ui(self):
        """Configurar interfaz del popup de crear factura"""
        layout = QVBoxLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Título
        title = "Crear Nueva Factura" if not self.datos_factura else "Editar Factura"
        title_label = QLabel(title)
        title_label.setObjectName("title_label")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Crear dos columnas para mejor organización
        main_form_layout = QHBoxLayout()
        
        # Columna izquierda
        left_form = QFormLayout()
        
        # Campo Empresa
        self.empresa_edit = QLineEdit()
        self.empresa_edit.setObjectName("input_field")
        self.empresa_edit.setPlaceholderText("Nombre de la empresa")
        left_form.addRow("Empresa:", self.empresa_edit)
        
        # Campo CIF
        self.cif_edit = QLineEdit()
        self.cif_edit.setObjectName("input_field")
        self.cif_edit.setPlaceholderText("CIF de la empresa")
        left_form.addRow("CIF:", self.cif_edit)
        
        # Campo Identificador Especial
        self.identificador_edit = QLineEdit()
        self.identificador_edit.setObjectName("input_field")
        self.identificador_edit.setPlaceholderText("Identificador especial")
        left_form.addRow("Identificador:", self.identificador_edit)
        
        # Campo Importe
        self.importe_spinbox = QDoubleSpinBox()
        self.importe_spinbox.setObjectName("input_field")
        self.importe_spinbox.setRange(0.0, 999999.99)
        self.importe_spinbox.setDecimals(2)
        self.importe_spinbox.setSuffix(" €")
        left_form.addRow("Importe:", self.importe_spinbox)
        
        # Campo Categoría
        self.categoria_combo = QComboBox()
        self.categoria_combo.setObjectName("input_field")
        categorias = ["Agua", "Vegetal", "Limpieza", "Actuaciones mantenimiento", "Otras"]
        self.categoria_combo.addItems(categorias)
        left_form.addRow("Categoría:", self.categoria_combo)
        
        # Columna derecha
        right_form = QFormLayout()
        
        # Campo Localidad
        self.localidad_edit = QLineEdit()
        self.localidad_edit.setObjectName("input_field")
        self.localidad_edit.setPlaceholderText("Ingrese la localidad")
        right_form.addRow("Localidad:", self.localidad_edit)
        
        # Campo GPED
        self.gped_edit = QLineEdit()
        self.gped_edit.setObjectName("input_field")
        self.gped_edit.setPlaceholderText("Ingrese GPED")
        right_form.addRow("GPED:", self.gped_edit)
        
        # Campo Estado
        self.estado_combo = QComboBox()
        self.estado_combo.setObjectName("input_field")
        estados = ["Emitida", "Tramitada", "Pagada", "Con deficiencias"]
        self.estado_combo.addItems(estados)
        right_form.addRow("Estado:", self.estado_combo)
        
        # Campo Identificación AdmYCont
        self.identificacion_admycont_edit = QLineEdit()
        self.identificacion_admycont_edit.setObjectName("input_field")
        self.identificacion_admycont_edit.setPlaceholderText("ID Administración y Control")
        right_form.addRow("ID AdmYCont:", self.identificacion_admycont_edit)
        
        # Campo Fecha Validación
        self.fecha_validacion_edit = QDateEdit()
        self.fecha_validacion_edit.setObjectName("input_field")
        self.fecha_validacion_edit.setDate(QDate.currentDate())  # Fecha actual por defecto
        self.fecha_validacion_edit.setCalendarPopup(True)  # Calendario desplegable
        right_form.addRow("Fecha Validación:", self.fecha_validacion_edit)
        
        # Añadir columnas al layout principal
        main_form_layout.addLayout(left_form)
        main_form_layout.addLayout(right_form)
        layout.addLayout(main_form_layout)
        
        # Campo Comentarios (ancho completo)
        comentarios_layout = QFormLayout()
        self.comentarios_edit = QTextEdit()
        self.comentarios_edit.setObjectName("text_input_field")
        self.comentarios_edit.setPlaceholderText("Ingrese comentarios adicionales...")
        self.comentarios_edit.setMaximumHeight(80)  # Altura limitada
        comentarios_layout.addRow("Comentarios:", self.comentarios_edit)
        layout.addLayout(comentarios_layout)
        
        # Sección PDF
        pdf_group = QGroupBox("Archivo PDF")
        pdf_group.setObjectName("group_box")
        pdf_layout = QVBoxLayout(pdf_group)
        
        # Label para mostrar archivo seleccionado
        self.archivo_label = QLabel("No hay archivo seleccionado")
        self.archivo_label.setObjectName("info_label")
        pdf_layout.addWidget(self.archivo_label)
        
        # Botones PDF
        pdf_buttons_layout = QHBoxLayout()
        
        self.btn_seleccionar_pdf = QPushButton("Seleccionar PDF")
        self.btn_seleccionar_pdf.setObjectName("action_button")
        self.btn_seleccionar_pdf.clicked.connect(self.seleccionar_pdf)
        pdf_buttons_layout.addWidget(self.btn_seleccionar_pdf)
        
        self.btn_quitar_pdf = QPushButton("Quitar PDF")
        self.btn_quitar_pdf.setObjectName("action_button")
        self.btn_quitar_pdf.clicked.connect(self.quitar_pdf)
        self.btn_quitar_pdf.setEnabled(False)
        pdf_buttons_layout.addWidget(self.btn_quitar_pdf)
        
        pdf_layout.addLayout(pdf_buttons_layout)
        layout.addWidget(pdf_group)
        
        # Botones principales
        buttons_layout = QHBoxLayout()
        
        self.btn_guardar = QPushButton("Guardar")
        self.btn_guardar.setObjectName("save_button")
        self.btn_guardar.clicked.connect(self.guardar_factura)
        buttons_layout.addWidget(self.btn_guardar)
        
        self.btn_cancelar = QPushButton("Cancelar")
        self.btn_cancelar.setObjectName("cancel_button")
        self.btn_cancelar.clicked.connect(self.reject)
        buttons_layout.addWidget(self.btn_cancelar)
        
        layout.addLayout(buttons_layout)
    
    def apply_adif_style(self):
        """Aplicar estilo ADIF Verde Corporativo"""
        style = """
        QDialog {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #f8f9fa, stop:1 #e9ecef);
            border: 2px solid #2E7D32;
            border-radius: 12px;
            font-family: 'Segoe UI', Arial, sans-serif;
        }
        
        QLabel#title_label {
            font-size: 11pt;
            font-weight: bold;
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #2E7D32, stop:1 #4CAF50);
            color: white;
            padding: 8px 15px;
            border-radius: 6px;
            margin: 5px 0px;
        }
        
        QDoubleSpinBox#input_field, QComboBox#input_field, QLineEdit#input_field, QDateEdit#input_field {
            font-size: 7pt;
            padding: 6px;
            border: 2px solid #4CAF50;
            border-radius: 4px;
            background: white;
        }
        
        QDoubleSpinBox#input_field:focus, QComboBox#input_field:focus, QLineEdit#input_field:focus, QDateEdit#input_field:focus {
            border: 2px solid #2E7D32;
        }
        
        QTextEdit#text_input_field {
            font-size: 7pt;
            padding: 6px;
            border: 2px solid #4CAF50;
            border-radius: 4px;
            background: white;
        }
        
        QTextEdit#text_input_field:focus {
            border: 2px solid #2E7D32;
        }
        
        QGroupBox#group_box {
            font-size: 7pt;
            font-weight: bold;
            color: #1B5E20;
            border: 2px solid #4CAF50;
            border-radius: 6px;
            margin-top: 8px;
            padding-top: 8px;
        }
        
        QLabel#info_label {
            font-size: 7pt;
            color: #388E3C;
            padding: 4px;
        }
        
        QPushButton#action_button {
            font-size: 7pt;
            font-weight: 500;
            color: white;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4CAF50, stop:1 #2E7D32);
            border: 1px solid #1B5E20;
            border-radius: 4px;
            padding: 6px 12px;
            margin: 2px;
        }
        
        QPushButton#action_button:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #66BB6A, stop:1 #388E3C);
        }
        
        QPushButton#save_button {
            font-size: 7pt;
            font-weight: bold;
            color: white;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #2E7D32, stop:1 #1B5E20);
            border: 2px solid #1B5E20;
            border-radius: 6px;
            padding: 8px 16px;
        }
        
        QPushButton#save_button:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4CAF50, stop:1 #2E7D32);
        }
        
        QPushButton#cancel_button {
            font-size: 7pt;
            font-weight: 500;
            color: #1B5E20;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f5f5f5);
            border: 2px solid #2E7D32;
            border-radius: 6px;
            padding: 8px 16px;
        }
        
        QPushButton#cancel_button:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #e8f5e8, stop:1 #c8e6c9);
        }
        """
        self.setStyleSheet(style)
    
    def cargar_datos_existentes(self):
        """Cargar datos de factura existente para edición"""
        if not self.datos_factura:
            return
        
        # Cargar todos los campos
        self.empresa_edit.setText(self.datos_factura.get('empresa', ''))
        self.cif_edit.setText(self.datos_factura.get('cif', ''))
        self.identificador_edit.setText(self.datos_factura.get('identificador_especial', ''))
        self.importe_spinbox.setValue(self.datos_factura.get('importe', 0.0))
        
        # Cargar ComboBox Categoría
        categoria = self.datos_factura.get('categoria', '')
        index = self.categoria_combo.findText(categoria)
        if index >= 0:
            self.categoria_combo.setCurrentIndex(index)
        
        # Cargar ComboBox Estado
        estado = self.datos_factura.get('estado', 'Emitida')
        index = self.estado_combo.findText(estado)
        if index >= 0:
            self.estado_combo.setCurrentIndex(index)
        
        self.localidad_edit.setText(self.datos_factura.get('localidad', ''))
        self.gped_edit.setText(self.datos_factura.get('gped', ''))
        self.comentarios_edit.setPlainText(self.datos_factura.get('comentarios', ''))
        
        # Cargar nuevos campos
        self.identificacion_admycont_edit.setText(self.datos_factura.get('identificacion_admycont', ''))
        
        # Cargar fecha validación
        fecha_validacion = self.datos_factura.get('fecha_validacion', '')
        if fecha_validacion:
            try:
                fecha = datetime.strptime(fecha_validacion, "%Y-%m-%d").date()
                self.fecha_validacion_edit.setDate(QDate(fecha.year, fecha.month, fecha.day))
            except:
                self.fecha_validacion_edit.setDate(QDate.currentDate())
        
        archivo_pdf = self.datos_factura.get('archivo_pdf', '')
        if archivo_pdf:
            self.archivo_pdf_seleccionado = archivo_pdf
            self.archivo_label.setText(f"Archivo: {archivo_pdf}")
            self.btn_quitar_pdf.setEnabled(True)
    
    def seleccionar_pdf(self):
        """Seleccionar archivo PDF"""
        archivo, _ = QFileDialog.getOpenFileName(
            self, "Seleccionar archivo PDF", "", "Archivos PDF (*.pdf)"
        )
        
        if archivo:
            # Validar que los campos necesarios estén llenos
            if not self.empresa_edit.text().strip():
                QMessageBox.warning(self, "Campo requerido", 
                                  "Debe ingresar el nombre de la empresa antes de seleccionar el PDF")
                return
            
            if self.importe_spinbox.value() <= 0:
                QMessageBox.warning(self, "Importe requerido", 
                                  "Debe ingresar un importe válido antes de seleccionar el PDF")
                return
            
            # Preparar datos para generar nombre del PDF
            if self.datos_factura:
                # Para edición, usar ID existente
                id_temp = self.datos_factura.get("id", 0)
            else:
                # Para nueva factura, obtener próximo ID
                if self.controlador:
                    datos_json = self.controlador.leer_datos_json()
                    id_temp = datos_json.get("configuracion", {}).get("ultimo_id", 0) + 1
                else:
                    id_temp = 1
                    
            datos_temp = {
                "id": id_temp,
                "empresa": self.empresa_edit.text().strip(),
                "importe": self.importe_spinbox.value()
            }
            
            # Copiar archivo a la carpeta de facturas directas
            if self.controlador:
                nuevo_nombre = self.controlador.gestionar_pdf(archivo, datos_temp)
                if nuevo_nombre:
                    self.archivo_pdf_seleccionado = nuevo_nombre
                    self.archivo_label.setText(f"Archivo: {nuevo_nombre}")
                    self.btn_quitar_pdf.setEnabled(True)
                    logger.info(f"PDF seleccionado con nombre: {nuevo_nombre}")
                else:
                    QMessageBox.critical(self, "Error", "No se pudo copiar el archivo PDF")
            else:
                QMessageBox.critical(self, "Error", "Controlador no disponible")
    
    def quitar_pdf(self):
        """Quitar archivo PDF seleccionado"""
        if self.archivo_pdf_seleccionado:
            reply = QMessageBox.question(self, "Quitar PDF", 
                                       "¿Desea eliminar el archivo PDF de la carpeta?",
                                       QMessageBox.Yes | QMessageBox.No,
                                       QMessageBox.No)
            if reply == QMessageBox.Yes and self.controlador:
                self.controlador.eliminar_pdf(self.archivo_pdf_seleccionado)
        
        self.archivo_pdf_seleccionado = ""
        self.archivo_label.setText("No hay archivo seleccionado")
        self.btn_quitar_pdf.setEnabled(False)
    
    def guardar_factura(self):
        """Guardar la factura"""
        try:
            # Validar campos obligatorios
            if not self.empresa_edit.text().strip():
                QMessageBox.warning(self, "Campo requerido", "El nombre de la empresa es obligatorio")
                return
            
            if not self.localidad_edit.text().strip():
                QMessageBox.warning(self, "Campo requerido", "La localidad es obligatoria")
                return
            
            if self.importe_spinbox.value() <= 0:
                QMessageBox.warning(self, "Importe inválido", "El importe debe ser mayor a 0")
                return
            
            # Preparar datos completos con nuevos campos
            datos = {
                "empresa": self.empresa_edit.text().strip(),
                "cif": self.cif_edit.text().strip(),
                "identificador_especial": self.identificador_edit.text().strip(),
                "identificacion_admycont": self.identificacion_admycont_edit.text().strip(),
                "fecha_validacion": self.fecha_validacion_edit.date().toString("yyyy-MM-dd"),
                "importe": self.importe_spinbox.value(),
                "categoria": self.categoria_combo.currentText(),
                "estado": self.estado_combo.currentText(),
                "localidad": self.localidad_edit.text().strip(),
                "gped": self.gped_edit.text().strip(),
                "comentarios": self.comentarios_edit.toPlainText().strip(),
                "archivo_pdf": self.archivo_pdf_seleccionado
            }
            
            # Si hay PDF pero aún no tiene el nombre correcto, regenerarlo
            if self.archivo_pdf_seleccionado and not self.datos_factura:  # Solo para nuevas facturas
                # El PDF ya debería tener el nombre correcto por seleccionar_pdf()
                pass
            
            # Guardar
            if self.controlador:
                if self.datos_factura:  # Edición
                    # Para edición, necesitamos actualizar la factura existente
                    datos["id"] = self.datos_factura["id"]
                    # Aquí iría la lógica de actualización
                    QMessageBox.information(self, "Función pendiente", 
                                          "La edición de facturas será implementada próximamente")
                else:  # Nueva factura
                    if self.controlador.agregar_factura(datos):
                        self.accept()
                    else:
                        QMessageBox.critical(self, "Error", "No se pudo guardar la factura")
            else:
                QMessageBox.critical(self, "Error", "Controlador no disponible")
        
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error guardando factura: {str(e)}")


class DialogoTablaResumen(QDialog):
    """Diálogo para mostrar resumen de facturas en tabla"""
    
    def __init__(self, parent=None, facturas=None):
        super().__init__(parent)
        self.facturas = facturas or []
        self.setWindowTitle("Resumen de Facturación")
        self.setFixedSize(1400, 800)  # Más amplio para filtros y columnas
        self.setWindowFlags(Qt.Dialog | Qt.WindowCloseButtonHint)
        self.setup_ui()
        self.apply_adif_style()
        self.cargar_datos_tabla()
    
    def setup_ui(self):
        """Configurar interfaz del diálogo de tabla"""
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Título
        title_label = QLabel("📊 Resumen de Facturación")
        title_label.setObjectName("title_label")
        title_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(title_label)
        
        # Información general
        total = sum(f.get('importe', 0) for f in self.facturas)
        info_label = QLabel(f"Total facturas: {len(self.facturas)} | Importe total: {total:.2f} €")
        info_label.setObjectName("info_label")
        info_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(info_label)
        
        # Sistema de filtros
        filtros_group = QGroupBox("🔍 Filtros de Búsqueda")
        filtros_group.setObjectName("filter_group")
        filtros_layout = QGridLayout(filtros_group)
        
        # Fila 1 de filtros
        filtros_layout.addWidget(QLabel("Empresa:"), 0, 0)
        self.filtro_empresa = QLineEdit()
        self.filtro_empresa.setObjectName("filter_input")
        self.filtro_empresa.setPlaceholderText("Buscar por empresa...")
        self.filtro_empresa.textChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_empresa, 0, 1)
        
        filtros_layout.addWidget(QLabel("CIF:"), 0, 2)
        self.filtro_cif = QLineEdit()
        self.filtro_cif.setObjectName("filter_input")
        self.filtro_cif.setPlaceholderText("Buscar por CIF...")
        self.filtro_cif.textChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_cif, 0, 3)
        
        # Fila 2 de filtros
        filtros_layout.addWidget(QLabel("Estado:"), 1, 0)
        self.filtro_estado = QComboBox()
        self.filtro_estado.setObjectName("filter_input")
        self.filtro_estado.addItems(["Todos", "Emitida", "Tramitada", "Pagada", "Con deficiencias"])
        self.filtro_estado.currentTextChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_estado, 1, 1)
        
        filtros_layout.addWidget(QLabel("Categoría:"), 1, 2)
        self.filtro_categoria = QComboBox()
        self.filtro_categoria.setObjectName("filter_input")
        self.filtro_categoria.addItems(["Todas", "Agua", "Vegetal", "Limpieza", "Actuaciones mantenimiento", "Otras"])
        self.filtro_categoria.currentTextChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_categoria, 1, 3)
        
        # Fila 3 de filtros
        filtros_layout.addWidget(QLabel("Fecha Validación:"), 2, 0)
        self.filtro_fecha_desde = QDateEdit()
        self.filtro_fecha_desde.setObjectName("filter_input")
        self.filtro_fecha_desde.setDate(QDate.currentDate().addDays(-365))  # Hace un año
        self.filtro_fecha_desde.setCalendarPopup(True)
        self.filtro_fecha_desde.dateChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_fecha_desde, 2, 1)
        
        filtros_layout.addWidget(QLabel("hasta:"), 2, 2)
        self.filtro_fecha_hasta = QDateEdit()
        self.filtro_fecha_hasta.setObjectName("filter_input")
        self.filtro_fecha_hasta.setDate(QDate.currentDate().addDays(365))  # Próximo año
        self.filtro_fecha_hasta.setCalendarPopup(True)
        self.filtro_fecha_hasta.dateChanged.connect(self.aplicar_filtros)
        filtros_layout.addWidget(self.filtro_fecha_hasta, 2, 3)
        
        # Botón limpiar filtros
        btn_limpiar = QPushButton("🗑️ Limpiar Filtros")
        btn_limpiar.setObjectName("clear_filters_button")
        btn_limpiar.clicked.connect(self.limpiar_filtros)
        filtros_layout.addWidget(btn_limpiar, 3, 0, 1, 4)
        
        layout.addWidget(filtros_group)
        
        # Guardar todas las facturas originales para filtros
        self.todas_facturas = self.facturas.copy()
        
        # Tabla
        self.tabla = QTableWidget()
        self.tabla.setObjectName("tabla_facturas")
        
        # Definir columnas con todos los campos incluyendo nuevos
        columnas = ["ID", "Fecha", "Empresa", "CIF", "Importe", "Estado", "Categoría", "Localidad", "GPED", "ID Especial", "ID AdmYCont", "Fecha Valid.", "Comentarios", "PDF"]
        self.tabla.setColumnCount(len(columnas))
        self.tabla.setHorizontalHeaderLabels(columnas)
        
        # Configurar tabla
        self.tabla.setAlternatingRowColors(True)
        self.tabla.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.tabla.horizontalHeader().setStretchLastSection(True)
        
        # Conectar evento de doble clic para editar estado
        self.tabla.cellDoubleClicked.connect(self.editar_estado_celda)
        
        layout.addWidget(self.tabla)
        
        # Botón cerrar
        btn_cerrar = QPushButton("Cerrar")
        btn_cerrar.setObjectName("close_button")
        btn_cerrar.setFixedHeight(35)
        btn_cerrar.clicked.connect(self.close)
        layout.addWidget(btn_cerrar)
    
    def cargar_datos_tabla(self):
        """Cargar datos de facturas en la tabla con todas las columnas"""
        self.tabla.setRowCount(len(self.facturas))
        
        for row, factura in enumerate(self.facturas):
            col = 0
            
            # ID
            self.tabla.setItem(row, col, QTableWidgetItem(str(factura.get('id', ''))))
            col += 1
            
            # Fecha (solo fecha, sin hora)
            fecha_completa = factura.get('fecha_creacion', '')
            fecha_corta = fecha_completa.split('T')[0] if 'T' in fecha_completa else fecha_completa
            self.tabla.setItem(row, col, QTableWidgetItem(fecha_corta))
            col += 1
            
            # Empresa
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('empresa', '')))
            col += 1
            
            # CIF
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('cif', '')))
            col += 1
            
            # Importe
            importe = factura.get('importe', 0)
            self.tabla.setItem(row, col, QTableWidgetItem(f"{importe:.2f} €"))
            col += 1
            
            # Estado
            estado = factura.get('estado', 'Emitida')
            item_estado = QTableWidgetItem(estado)
            # Colorear según estado
            if estado == "Pagada":
                item_estado.setBackground(QColor(200, 255, 200))  # Verde claro
            elif estado == "Con deficiencias":
                item_estado.setBackground(QColor(255, 200, 200))  # Rojo claro
            elif estado == "Tramitada":
                item_estado.setBackground(QColor(255, 255, 200))  # Amarillo claro
            self.tabla.setItem(row, col, item_estado)
            col += 1
            
            # Categoría
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('categoria', '')))
            col += 1
            
            # Localidad
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('localidad', '')))
            col += 1
            
            # GPED
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('gped', '')))
            col += 1
            
            # Identificador Especial
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('identificador_especial', '')))
            col += 1
            
            # ID AdmYCont (nuevo campo)
            self.tabla.setItem(row, col, QTableWidgetItem(factura.get('identificacion_admycont', '')))
            col += 1
            
            # Fecha Validación (nuevo campo)
            fecha_validacion = factura.get('fecha_validacion', '')
            self.tabla.setItem(row, col, QTableWidgetItem(fecha_validacion))
            col += 1
            
            # Comentarios (truncar si es muy largo)
            comentarios = factura.get('comentarios', '')
            comentarios_cortos = comentarios[:25] + "..." if len(comentarios) > 25 else comentarios
            self.tabla.setItem(row, col, QTableWidgetItem(comentarios_cortos))
            col += 1
            
            # PDF
            pdf = factura.get('archivo_pdf', '')
            pdf_nombre = "✅" if pdf else "❌"
            self.tabla.setItem(row, col, QTableWidgetItem(pdf_nombre))
        
        # Ajustar tamaño de columnas
        self.tabla.resizeColumnsToContents()
    
    def editar_estado_celda(self, row, col):
        """Editar estado de factura con doble clic en columna Estado"""
        try:
            # La columna Estado es la número 5 (índice basado en 0)
            if col == 5:  # Columna Estado
                factura_id = int(self.tabla.item(row, 0).text())  # ID en columna 0
                estado_actual = self.tabla.item(row, col).text()
                
                # Mostrar diálogo de selección de estado
                estados = ["Emitida", "Tramitada", "Pagada", "Con deficiencias"]
                nuevo_estado, ok = QInputDialog.getItem(
                    self, f"Editar Estado - Factura {factura_id}", 
                    "Seleccione el nuevo estado:", 
                    estados, estados.index(estado_actual), False
                )
                
                if ok and nuevo_estado != estado_actual:
                    # Actualizar en el controlador
                    if hasattr(self.parent(), 'controlador') and self.parent().controlador:
                        controlador = self.parent().controlador
                        if controlador.actualizar_estado_factura(factura_id, nuevo_estado):
                            # Actualizar visualmente en la tabla
                            item_estado = QTableWidgetItem(nuevo_estado)
                            
                            # Colorear según estado
                            if nuevo_estado == "Pagada":
                                item_estado.setBackground(QColor(200, 255, 200))
                            elif nuevo_estado == "Con deficiencias":
                                item_estado.setBackground(QColor(255, 200, 200))
                            elif nuevo_estado == "Tramitada":
                                item_estado.setBackground(QColor(255, 255, 200))
                            
                            self.tabla.setItem(row, col, item_estado)
                            QMessageBox.information(self, "Estado actualizado", 
                                                  f"Estado de factura {factura_id} cambiado a: {nuevo_estado}")
                        else:
                            QMessageBox.critical(self, "Error", "No se pudo actualizar el estado")
                    else:
                        QMessageBox.critical(self, "Error", "Controlador no disponible")
            else:
                # Mostrar información de la celda para otras columnas
                item = self.tabla.item(row, col)
                if item and item.text():
                    QMessageBox.information(self, "Información", 
                                          f"Columna: {self.tabla.horizontalHeaderItem(col).text()}\n"
                                          f"Valor: {item.text()}")
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error editando estado: {str(e)}")
    
    def aplicar_filtros(self):
        """Aplicar filtros a la tabla de facturas"""
        try:
            # Obtener valores de los filtros
            empresa_filtro = self.filtro_empresa.text().lower().strip()
            cif_filtro = self.filtro_cif.text().lower().strip()
            estado_filtro = self.filtro_estado.currentText()
            categoria_filtro = self.filtro_categoria.currentText()
            fecha_desde = self.filtro_fecha_desde.date().toPyDate()
            fecha_hasta = self.filtro_fecha_hasta.date().toPyDate()
            
            # Filtrar facturas
            facturas_filtradas = []
            
            for factura in self.todas_facturas:
                # Filtro por empresa
                if empresa_filtro and empresa_filtro not in factura.get('empresa', '').lower():
                    continue
                    
                # Filtro por CIF
                if cif_filtro and cif_filtro not in factura.get('cif', '').lower():
                    continue
                    
                # Filtro por estado
                if estado_filtro != "Todos" and factura.get('estado', '') != estado_filtro:
                    continue
                    
                # Filtro por categoría
                if categoria_filtro != "Todas" and factura.get('categoria', '') != categoria_filtro:
                    continue
                    
                # Filtro por fecha de validación
                fecha_validacion_str = factura.get('fecha_validacion', '')
                if fecha_validacion_str:
                    try:
                        fecha_validacion = datetime.strptime(fecha_validacion_str, "%Y-%m-%d").date()
                        if fecha_validacion < fecha_desde or fecha_validacion > fecha_hasta:
                            continue
                    except:
                        # Si no se puede parsear la fecha, incluir la factura
                        pass
                
                # Si llegamos aquí, la factura pasa todos los filtros
                facturas_filtradas.append(factura)
            
            # Actualizar tabla con facturas filtradas
            self.facturas = facturas_filtradas
            self.cargar_datos_tabla()
            
            # Actualizar información general
            total = sum(f.get('importe', 0) for f in facturas_filtradas)
            info_text = f"Total facturas mostradas: {len(facturas_filtradas)} de {len(self.todas_facturas)} | Importe total: {total:.2f} €"
            
            # Buscar el label de información y actualizarlo
            for child in self.children():
                if isinstance(child, QLabel) and hasattr(child, 'objectName') and child.objectName() == "info_label":
                    child.setText(info_text)
                    break
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error aplicando filtros: {str(e)}")
    
    def limpiar_filtros(self):
        """Limpiar todos los filtros y mostrar todas las facturas"""
        try:
            self.filtro_empresa.clear()
            self.filtro_cif.clear()
            self.filtro_estado.setCurrentText("Todos")
            self.filtro_categoria.setCurrentText("Todas")
            self.filtro_fecha_desde.setDate(QDate.currentDate().addDays(-365))
            self.filtro_fecha_hasta.setDate(QDate.currentDate().addDays(365))
            
            # Restaurar todas las facturas
            self.facturas = self.todas_facturas.copy()
            self.cargar_datos_tabla()
            
            # Actualizar información general
            total = sum(f.get('importe', 0) for f in self.facturas)
            info_text = f"Total facturas: {len(self.facturas)} | Importe total: {total:.2f} €"
            
            # Buscar el label de información y actualizarlo
            for child in self.children():
                if isinstance(child, QLabel) and hasattr(child, 'objectName') and child.objectName() == "info_label":
                    child.setText(info_text)
                    break
                    
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error limpiando filtros: {str(e)}")
    
    def apply_adif_style(self):
        """Aplicar estilo ADIF Verde Corporativo a la tabla"""
        style = """
        QDialog {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #f8f9fa, stop:1 #e9ecef);
            border: 2px solid #2E7D32;
            border-radius: 12px;
            font-family: 'Segoe UI', Arial, sans-serif;
        }
        
        QLabel#title_label {
            font-size: 11pt;
            font-weight: bold;
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #2E7D32, stop:1 #4CAF50);
            color: white;
            padding: 8px 15px;
            border-radius: 6px;
            margin: 5px 0px;
        }
        
        QLabel#info_label {
            font-size: 8pt;
            font-weight: bold;
            color: #1B5E20;
            padding: 5px;
            background: rgba(76, 175, 80, 0.1);
            border-radius: 4px;
        }
        
        QTableWidget#tabla_facturas {
            font-size: 7pt;
            gridline-color: #4CAF50;
            background-color: white;
            alternate-background-color: rgba(76, 175, 80, 0.1);
            selection-background-color: #66BB6A;
            border: 1px solid #2E7D32;
            border-radius: 4px;
        }
        
        QTableWidget#tabla_facturas::item {
            padding: 4px;
            border: none;
        }
        
        QTableWidget#tabla_facturas::item:selected {
            background-color: #4CAF50;
            color: white;
        }
        
        QHeaderView::section {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4CAF50, stop:1 #2E7D32);
            color: white;
            padding: 6px;
            border: 1px solid #2E7D32;
            font-weight: bold;
            font-size: 7pt;
        }
        
        QPushButton#close_button {
            font-size: 7pt;
            font-weight: 500;
            color: #1B5E20;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f5f5f5);
            border: 2px solid #2E7D32;
            border-radius: 6px;
            padding: 6px 12px;
            margin: 2px;
        }
        
        QPushButton#close_button:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #e8f5e8, stop:1 #c8e6c9);
            border: 2px solid #4CAF50;
        }
        
        /* Estilos para filtros */
        QGroupBox#filter_group {
            font-size: 8pt;
            font-weight: bold;
            color: #1B5E20;
            border: 2px solid #4CAF50;
            border-radius: 8px;
            margin-top: 8px;
            padding-top: 8px;
            background: rgba(76, 175, 80, 0.05);
        }
        
        QLineEdit#filter_input, QComboBox#filter_input, QDateEdit#filter_input {
            font-size: 7pt;
            padding: 4px;
            border: 1px solid #4CAF50;
            border-radius: 3px;
            background: white;
        }
        
        QLineEdit#filter_input:focus, QComboBox#filter_input:focus, QDateEdit#filter_input:focus {
            border: 2px solid #2E7D32;
        }
        
        QPushButton#clear_filters_button {
            font-size: 7pt;
            font-weight: 500;
            color: white;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #FF7043, stop:1 #D84315);
            border: 1px solid #BF360C;
            border-radius: 4px;
            padding: 6px 12px;
            margin: 2px;
        }
        
        QPushButton#clear_filters_button:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #FF8A65, stop:1 #FF5722);
        }
        """
        self.setStyleSheet(style)
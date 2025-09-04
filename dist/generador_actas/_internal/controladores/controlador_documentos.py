"""
Controlador para generación de documentos Word con sustitución de variables
Versión mínima optimizada - Solo funcionalidades esenciales
"""
import os
import re
import sys, json
import subprocess
import logging
from typing import Dict, Any, List, Optional

logger = logging.getLogger(__name__)
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PyQt5.QtWidgets import QMessageBox

from helpers_py import (
    resource_path, limpiar_nombre_archivo, formatear_numero_espanol
)
# Importación condicional del módulo PDF para evitar errores
try:
    from .controlador_pdf_unificado import mostrar_dialogo_pdf
    PDF_DISPONIBLE = True
except Exception as e:
    logger.warning(f"[ControladorDocumentos] Módulo PDF no disponible: {e}")
    PDF_DISPONIBLE = False
    mostrar_dialogo_pdf = None


class ControladorDocumentos:
    """Controlador para generación de documentos con sustitución de variables"""
    
    def __init__(self, main_window=None):
        """Inicializar controlador de documentos - CON GESTOR UNIFICADO"""
        self.main_window = main_window
        self.directorio_plantillas = "plantillas"
        
        # 🆕 NUEVO: Inicializar gestor de archivos
        self.gestor_archivos = None
        
        # Mapeo de funciones a plantillas base (existente)
        self.mapeo_plantillas_base = {
            'generar_acta_inicio': 'plantilla_acta_inicio',
            'generar_cartas_invitacion': 'plantilla_cartas_invitacion',
            'generar_acta_adjudicacion': 'plantilla_acta_adjudicacion',
            'generar_cartas_adjudicacion': 'plantilla_cartas_adjudicacion',
            'generar_acta_liquidacion': 'plantilla_acta_finalizacion',
            'generar_acta_replanteo': 'plantilla_acta_replanteo',
            'generar_acta_recepcion': 'plantilla_acta_recepcion',
            'generar_nombramiento_director': 'Modelo_director_obra_ajuste',
            'generar_contrato': 'plantilla_contrato'
        }
        
        # Mapeo legacy para compatibilidad (mantenido para no romper el código existente)
        self.mapeo_plantillas = {
            'generar_acta_inicio': 'plantilla_acta_inicio.docx',
            'generar_cartas_invitacion': 'plantilla_cartas_invitacion.docx',
            'generar_acta_adjudicacion': 'plantilla_acta_adjudicacion.docx',
            'generar_cartas_adjudicacion': 'plantilla_cartas_adjudicacion.docx',
            'generar_acta_liquidacion': 'plantilla_acta_finalizacion.docx',
            'generar_acta_replanteo': 'plantilla_acta_replanteo.docx',
            'generar_acta_recepcion': 'plantilla_acta_recepcion.docx',
            'generar_nombramiento_director': 'Modelo_director_obra_ajuste.docx',
            'generar_contrato': 'plantilla_contrato.docx'
        }
        
        # 🆕 NUEVO: Configurar gestor unificado
        self._configurar_gestor_unificado()
        
        # 🆕 NUEVO: Configurar tracker de documentos
        self._configurar_tracker_documentos()

    # ===== NUEVA FUNCIÓN A AÑADIR =====
    def _configurar_gestor_unificado(self):
        """Configurar gestor de archivos unificado"""
        try:
            if self.main_window and hasattr(self.main_window, 'gestor_archivos_unificado'):
                self.gestor_archivos = self.main_window.gestor_archivos_unificado
            else:
                logger.warning(f"[ControladorDocumentos] [!] Gestor unificado no disponible")
        except Exception as e:
            logger.error(f"[ControladorDocumentos] [X] Error configurando gestor: {e}")
    
    def _configurar_tracker_documentos(self):
        """Configurar tracker de documentos generados"""
        try:
            from .controlador_resumen import TrackerDocumentos, TipoDocumento
            self.tracker = TrackerDocumentos()
            self.TipoDocumento = TipoDocumento
            logger.info(f"[ControladorDocumentos] ✅ Tracker de documentos configurado")
            return True
        except Exception as e:
            logger.error(f"[ControladorDocumentos] [!] Error configurando tracker: {e}")
            self.tracker = None
            return False
    
    # =================== MÉTODOS REQUERIDOS POR LOS TESTS ===================
    
    def generar_documento_word(self, tipo_documento, nombre_archivo, datos_contrato, plantilla_nombre):
        """Generar documento Word con datos del contrato"""
        try:
            return self._generar_documento_con_sustitucion(tipo_documento, datos_contrato, nombre_archivo)
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error generando documento: {e}")
            return False
    
    def _sustituir_variables_texto(self, texto: str, datos_json: Dict[str, Any]) -> str:
        """Sustituir variables en texto usando formato @variable@"""
        try:
            # Buscar todas las variables en formato @variable@
            patron = r'@(\w+)@'
            
            def reemplazar_variable(match):
                nombre_variable = match.group(1)
                if nombre_variable in datos_json:
                    valor = datos_json[nombre_variable]
                    if isinstance(valor, (int, float)):
                        return formatear_numero_espanol(valor)
                    return str(valor)
                return match.group(0)  # Mantener la variable si no se encuentra
            
            return re.sub(patron, reemplazar_variable, texto)
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error sustituyendo variables: {e}")
            return texto
    
    def _validar_datos_basicos_contrato(self, datos_contrato: Dict[str, Any]) -> bool:
        """Validar que los datos básicos del contrato están completos"""
        campos_requeridos = [
            'nombreObra', 'numeroExpediente', 'basePresupuesto', 
            'precioAdjudicacion', 'empresaAdjudicada'
        ]
        
        for campo in campos_requeridos:
            if campo not in datos_contrato or not datos_contrato[campo]:
                return False
        return True
    
    def _validar_plantilla_existe(self, nombre_plantilla: str) -> bool:
        """Validar que una plantilla existe"""
        try:
            ruta_plantilla = self._obtener_ruta_plantilla(nombre_plantilla)
            return ruta_plantilla is not None and os.path.exists(ruta_plantilla)
        except Exception:
            return False
    
    def _obtener_nombre_carpeta_actual(self, nombre_contrato: str) -> str:
        """Obtener el nombre real de la carpeta desde el JSON (considerando cambios de expediente)"""
        try:
            # Obtener datos del contrato desde el JSON
            if hasattr(self.main_window, 'controlador_json') and self.main_window.controlador_json:
                contrato_data = self.main_window.controlador_json.leer_contrato_completo(nombre_contrato)
                if contrato_data and 'nombreCarpeta' in contrato_data:
                    nombre_carpeta = contrato_data['nombreCarpeta']
                    logger.info(f"[ControladorDocumentos] 📁 Nombre de carpeta desde JSON: {nombre_carpeta}")
                    return nombre_carpeta
            
            # Fallback: usar nombre del contrato si no hay nombreCarpeta en JSON
            logger.warning(f"[ControladorDocumentos] ⚠️ Usando nombre de contrato como fallback: {nombre_contrato}")
            return nombre_contrato
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error obteniendo nombre carpeta: {e}")
            return nombre_contrato  # Fallback seguro
    
    def _iniciar_tracking_documento(self, tipo_documento, nombre_documento, plantilla=""):
        """Iniciar tracking de un documento que se va a generar"""
        if not self.tracker or not hasattr(self, 'contract_name'):
            return None
        
        try:
            return self.tracker.registrar_documento_iniciado(
                self.contract_name,
                tipo_documento,
                nombre_documento,
                plantilla
            )
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error iniciando tracking: {e}")
            return None
    
    def _completar_tracking_documento(self, documento_id, ruta_archivo, observaciones=""):
        """Completar tracking de un documento generado exitosamente"""
        logger.debug(f"[ControladorDocumentos] 🔍 Completando tracking - ID: {documento_id}")
        logger.debug(f"[ControladorDocumentos] 🔍 Tracker: {self.tracker is not None}")
        logger.debug(f"[ControladorDocumentos] 🔍 Contract_name: {hasattr(self, 'contract_name')}")
        
        if not self.tracker or not documento_id or not hasattr(self, 'contract_name'):
            logger.warning(f"[ControladorDocumentos] ⚠️ Tracking no completado - faltan requisitos")
            return
        
        try:
            self.tracker.registrar_documento_completado(
                self.contract_name,
                documento_id,
                ruta_archivo,
                observaciones
            )
            
            # ✅ NOTA: La fase ya se actualizó durante la generación
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error completando tracking: {e}")
    
    def _actualizar_fase_en_generacion(self, tipo_documento: str):
        """Actualizar la fase del documento DURANTE la generación"""
        try:
            logger.info(f"[ControladorDocumentos] 🎯 Actualizando fase durante generación: {tipo_documento}")
            
            # Verificar disponibilidad del controlador de fases
            if not self.main_window or not hasattr(self.main_window, 'controlador_fases'):
                logger.warning(f"[ControladorDocumentos] ⚠️ Controlador de fases no disponible")
                return
            
            # Obtener nombre del contrato actual
            nombre_contrato = None
            if hasattr(self, 'contract_name'):
                nombre_contrato = self.contract_name
            elif hasattr(self.main_window, 'comboBox'):
                nombre_contrato = self.main_window.comboBox.currentText()
            
            if not nombre_contrato:
                logger.warning(f"[ControladorDocumentos] ⚠️ No se pudo determinar el contrato actual")
                return
            
            # Marcar documento como generado en el controlador de fases
            self.main_window.controlador_fases.marcar_documento_generado(tipo_documento, nombre_contrato)
            logger.info(f"[ControladorDocumentos] ✅ Fase actualizada durante generación: {tipo_documento}")
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error actualizando fase durante generación: {e}")
            import traceback
            logger.exception("Error completo:")
    
    def _notificar_documento_generado(self, documento_id: str, ruta_archivo: str):
        """Notificar al controlador de fases que se generó un documento"""
        try:
            logger.info(f"[ControladorDocumentos] 🔔 Notificando documento generado: {documento_id}")
            
            # Determinar tipo de documento por el ID o la ruta
            tipo_documento = self._detectar_tipo_documento(documento_id, ruta_archivo)
            logger.debug(f"[ControladorDocumentos] 📋 Tipo detectado: {tipo_documento}")
            
            # Verificar disponibilidad del controlador de fases
            logger.debug(f"[ControladorDocumentos] 🔍 Main window: {self.main_window is not None}")
            logger.debug(f"[ControladorDocumentos] 🔍 Tiene controlador_fases: {hasattr(self.main_window, 'controlador_fases') if self.main_window else False}")
            
            # Notificar al controlador de fases si está disponible
            if self.main_window and hasattr(self.main_window, 'controlador_fases'):
                logger.info(f"[ControladorDocumentos] ➡️ Enviando a controlador de fases: {tipo_documento} para {self.contract_name}")
                self.main_window.controlador_fases.marcar_documento_generado(
                    tipo_documento, self.contract_name
                )
                logger.info(f"[ControladorDocumentos] ✅ Fase actualizada para documento: {tipo_documento}")
            else:
                logger.warning(f"[ControladorDocumentos] ❌ Controlador de fases no disponible")
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ⚠️ Error notificando fases: {e}")
            import traceback
            logger.exception("Error completo:")
    
    def _detectar_tipo_documento(self, documento_id: str, ruta_archivo: str) -> str:
        """Detectar el tipo de documento basado en ID o ruta"""
        logger.debug(f"[ControladorDocumentos] 🔍 Detectando tipo - ID: {documento_id}, Ruta: {ruta_archivo}")
        
        # Convertir a minúsculas para comparación
        texto_busqueda = f"{documento_id} {ruta_archivo}".lower()
        logger.debug(f"[ControladorDocumentos] 🔍 Texto de búsqueda: {texto_busqueda}")
        
        # Mapeo mejorado de términos a tipos de documento
        mapeo_tipos = {
            "acta_inicio": ["inicio", "acta_inicio", "plantilla_acta_inicio"],
            "invitacion": ["invitacion", "cartas_invitacion", "carta", "plantilla_cartas_invitacion"],
            "adjudicacion": ["adjudicacion", "acta_adjudicacion", "plantilla_acta_adjudicacion"],
            "contrato": ["contrato", "firma_contrato", "plantilla_contrato"],
            "acta_replanteo": ["replanteo", "acta_replanteo", "plantilla_acta_replanteo"],
            "actuacion": ["actuacion", "ejecucion"],
            "acta_recepcion": ["recepcion", "acta_recepcion", "plantilla_acta_recepcion"],
            "acta_finalizacion": ["finalizacion", "liquidacion", "acta_finalizacion", "plantilla_acta_finalizacion"]
        }
        
        # Buscar coincidencias
        for tipo, palabras_clave in mapeo_tipos.items():
            for palabra in palabras_clave:
                if palabra in texto_busqueda:
                    logger.debug(f"[ControladorDocumentos] ✅ Tipo detectado: {tipo} (coincidencia: {palabra})")
                    return tipo
        
        # Si no se encuentra, intentar extraer del nombre del archivo
        nombre_archivo = os.path.basename(ruta_archivo).lower() if ruta_archivo else ""
        logger.debug(f"[ControladorDocumentos] 🔍 Verificando nombre archivo: {nombre_archivo}")
        
        for tipo, palabras_clave in mapeo_tipos.items():
            for palabra in palabras_clave:
                if palabra in nombre_archivo:
                    logger.debug(f"[ControladorDocumentos] ✅ Tipo detectado por archivo: {tipo} (coincidencia: {palabra})")
                    return tipo
        
        # También verificar el tracker si está disponible para obtener contexto
        if hasattr(self, 'contract_name') and hasattr(self, 'tracker'):
            try:
                # Obtener información del documento desde el tracker
                if self.tracker:
                    documentos = self.tracker.obtener_documentos_contrato(self.contract_name)
                    for doc in documentos:
                        if doc.id == documento_id:
                            logger.debug(f"[ControladorDocumentos] 🔍 Info del tracker: {doc.tipo.value if hasattr(doc.tipo, 'value') else doc.tipo}")
                            return str(doc.tipo.value if hasattr(doc.tipo, 'value') else doc.tipo)
            except Exception as e:
                logger.warning(f"[ControladorDocumentos] ⚠️ Error obteniendo info del tracker: {e}")
        
        # Fallback: devolver tipo genérico
        logger.warning(f"[ControladorDocumentos] ⚠️ No se pudo detectar tipo específico, usando 'documento'")
        return "documento"
    
    def _error_tracking_documento(self, documento_id, error):
        """Registrar error en tracking de documento"""
        if not self.tracker or not documento_id or not hasattr(self, 'contract_name'):
            return
        
        try:
            self.tracker.registrar_documento_error(
                self.contract_name,
                documento_id,
                str(error)
            )
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error registrando error en tracking: {e}")
    
    def _mapear_tipo_documento(self, tipo_funcion: str):
        """Mapear tipo de función a tipo de documento para tracking"""
        mapeo = {
            'generar_cartas_invitacion': 'cartas_invitacion',
            'generar_cartas_adjudicacion': 'adjudicacion', 
            'generar_acta_inicio': 'acta_inicio',
            'generar_acta_replanteo': 'acta_replanteo',
            'generar_acta_recepcion': 'acta_recepcion', 
            'generar_acta_liquidacion': 'liquidacion',
            'generar_contrato': 'contrato',
            'generar_acta_adjudicacion': 'adjudicacion',
            'generar_nombramiento_director': 'nombramiento'
        }
        
        return mapeo.get(tipo_funcion, 'otro')
    
    def _generar_documento_con_sustitucion_tracked(self, tipo_funcion: str, contract_data: Dict[str, Any], nombre_archivo: str):
        """Versión del método de sustitución que retorna también la ruta del archivo generado"""
        try:
            # Construir la ruta donde se generará el archivo antes de la generación
            carpeta_contrato = self._obtener_carpeta_con_gestor_unificado(contract_data)
            if not carpeta_contrato:
                return False, ""
                
            subcarpeta_destino = self._determinar_subcarpeta_por_tipo_documento(tipo_funcion)
            nombre_archivo_limpio = self._limpiar_nombre_archivo(f"{nombre_archivo}")
            archivo_salida = os.path.join(carpeta_contrato, subcarpeta_destino, f"{nombre_archivo_limpio}.docx")
            
            # Llamar al método original
            exito = self._generar_documento_con_sustitucion(tipo_funcion, contract_data, nombre_archivo)
            
            if exito:
                # Verificar que el archivo se haya generado correctamente
                if os.path.exists(archivo_salida):
                    logger.info(f"[ControladorDocumentos] ✅ Archivo generado: {archivo_salida}")
                    return True, archivo_salida
                else:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Archivo esperado no encontrado: {archivo_salida}")
                    return True, ""
            else:
                return False, ""
                
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error en generación tracked: {e}")
            return False, ""
    
    def _obtener_carpeta_contrato_actual(self):
        """Obtener la carpeta del contrato actual"""
        try:
            if hasattr(self, 'gestor_archivos') and self.gestor_archivos:
                contract_data = self._obtener_datos_contrato_actual()
                if contract_data:
                    return self.gestor_archivos.obtener_carpeta_obra(contract_data)
            
            # Fallback: usar método legacy si está disponible
            if hasattr(self, 'main_window') and self.main_window:
                return getattr(self.main_window, 'current_work_folder', None)
            
            return None
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error obteniendo carpeta contrato: {e}")
            return None

    # ===== NUEVA FUNCIÓN A AÑADIR =====
    def _obtener_carpeta_con_gestor_unificado(self, contract_data):
        """Obtener carpeta usando el gestor unificado"""
        try:
            
            # Usar gestor configurado
            if self.gestor_archivos:
                carpeta_path = self.gestor_archivos.obtener_carpeta_obra(contract_data, crear_si_no_existe=True)
                
                if carpeta_path:
                    return carpeta_path
                else:
                    logger.error(f"[ControladorDocumentos] ❌ Gestor no retornó carpeta")
            else:
                logger.warning(f"[ControladorDocumentos] ⚠️ Gestor no disponible")
            
            # Fallback al método anterior
            return self._crear_o_obtener_carpeta_contrato(contract_data)
                
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error obteniendo carpeta: {e}")
            return self._crear_o_obtener_carpeta_contrato(contract_data)

    # ===== NUEVA FUNCIÓN A AÑADIR =====
    def _determinar_subcarpeta_por_tipo_documento(self, tipo_funcion):
        """Determinar subcarpeta según tipo de documento"""
        try:
            # MAPEO COMPLETO DE TIPOS A SUBCARPETAS
            mapeo_subcarpetas = {
                # ACTAS - Van a documentos finales (firmados)
                'generar_acta_inicio': '04-documentos-sin-firmar',
                'generar_acta_adjudicacion': '04-documentos-sin-firmar', 
                'generar_acta_liquidacion': '04-documentos-sin-firmar',
                'generar_acta_replanteo': '04-documentos-sin-firmar',
                'generar_acta_recepcion': '04-documentos-sin-firmar',
                
                # CARTAS - Diferentes destinos según tipo
                'generar_cartas_invitacion': '05-cartas-sin-firmar',
                'generar_cartas_adjudicacion': '03-cartas-finales',
                
                # NOMBRAMIENTOS - Van a documentación final
                'generar_nombramiento_director': '04-documentos-sin-firmar',
                
                # CONTRATOS - Van a documentos finales
                'generar_contrato': '04-documentos-sin-firmar',
                
                # DOCUMENTOS GENERALES
                'documento_generico': '04-documentos-sin-firmar',
                'informe': '10-otros',
                'oferta': '06-ofertas'
            }
            
            subcarpeta = mapeo_subcarpetas.get(tipo_funcion, '04-documentos-sin-firmar')
            return subcarpeta
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error determinando subcarpeta: {e}")
            return '04-documentos-sin-firmar'

    
    def set_main_window(self, main_window):
        """Establecer referencia a la ventana principal"""
        self.main_window = main_window

    def _obtener_nombre_plantilla_dinamico(self, tipo_funcion: str) -> str:
        """Obtener nombre de plantilla dinámico según tipo de contrato"""
        try:
            # Obtener nombre base de la plantilla
            plantilla_base = self.mapeo_plantillas_base.get(tipo_funcion)
            
            if not plantilla_base:
                # Fallback al mapeo legacy
                plantilla_legacy = self.mapeo_plantillas.get(tipo_funcion)
                return plantilla_legacy if plantilla_legacy else f"{tipo_funcion}.docx"
            
            # Determinar tipo basado en el nuevo sistema
            es_obra = False
            if (self.main_window and 
                hasattr(self.main_window, 'contract_manager') and 
                self.main_window.contract_manager):
                
                contract_data = self.main_window.contract_manager.get_current_contract_data()
                if contract_data:
                    tipo_actuacion = contract_data.get('tipoActuacion', '')
                    es_obra = tipo_actuacion in ['obras', 'obra_mantenimiento']
            
            # Construir nombre de archivo según tipo
            if es_obra:
                nombre_plantilla = f"{plantilla_base}_obra.docx"
                logger.debug(f"[ControladorDocumentos] [OBRA] Tipo obra/obra_mantenimiento - Usando plantilla OBRA: {nombre_plantilla}")
            else:
                nombre_plantilla = f"{plantilla_base}_servicio.docx"
                logger.debug(f"[ControladorDocumentos] [SERVICIO] Tipo servicio/serv_mantenimiento - Usando plantilla SERVICIO: {nombre_plantilla}")
            
            return nombre_plantilla
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] [X] Error obteniendo plantilla dinámica: {e}")
            # Fallback al mapeo legacy
            return self.mapeo_plantillas.get(tipo_funcion, f"{tipo_funcion}.docx")

    # =================== MÉTODOS DE GENERACIÓN ===================

    def generar_acta_inicio(self, datos_contrato=None, nombre_archivo=None):
        """Generar acta de inicio CON selección dinámica de plantilla"""
        try:
            
            # 🆕 NUEVO: Verificación de plantilla dinámica
            nombre_plantilla = self._obtener_nombre_plantilla_dinamico('generar_acta_inicio')
            if not self.verificar_plantilla_disponible(nombre_plantilla):
                # Fallback a plantilla legacy
                nombre_plantilla_legacy = 'plantilla_acta_inicio.docx'
                if not self.verificar_plantilla_disponible(nombre_plantilla_legacy):
                    self._mostrar_error(f"No se encontraron plantillas: {nombre_plantilla} ni {nombre_plantilla_legacy}")
                    return
            
            # Continuar con la generación normal
            self._generar_documento_simple('generar_acta_inicio', 'Acta_Inicio', "Acta de Inicio")
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error en generar_acta_inicio: {e}")
            import traceback
            logger.exception("Error completo:")
            self._mostrar_error(f"Error inesperado generando Acta de Inicio: {str(e)}")

    def generar_cartas_invitacion(self, datos_contrato=None, empresas_lista=None):
        """Generar cartas de invitación individuales - CON GESTOR UNIFICADO"""
        try:
            if datos_contrato is None and not self._validar_contrato_seleccionado():
                return
            
            contract_data = self._obtener_datos_contrato_actual()
            if not contract_data:
                return self._mostrar_error("No se pudieron obtener los datos del contrato")
            
            empresas_lista = self._obtener_empresas_lista(contract_data)
            if not empresas_lista:
                return self._mostrar_error("No hay empresas para generar cartas de invitación")
            
            ruta_plantilla = self._obtener_ruta_plantilla('plantilla_cartas_invitacion.docx')
            if not ruta_plantilla:
                return self._mostrar_error("No se encontró la plantilla de cartas de invitación")
            
            # 🆕 USAR GESTOR UNIFICADO
            carpeta_contrato = self._obtener_carpeta_con_gestor_unificado(contract_data)
            if not carpeta_contrato:
                return self._mostrar_error("No se pudo crear/obtener la carpeta del contrato")
            
            # 🆕 USAR SUBCARPETA CORRECTA
            subcarpeta_cartas = self._determinar_subcarpeta_por_tipo_documento('generar_cartas_invitacion')
            directorio_cartas = os.path.join(carpeta_contrato, subcarpeta_cartas)
            os.makedirs(directorio_cartas, exist_ok=True)
            
            # Resto del código igual...
            cartas_generadas = []
            for i, empresa in enumerate(empresas_lista):
                nombre_empresa_limpio = self._limpiar_nombre_para_archivo(empresa.get('nombre', f'Empresa_{i+1}'))
                nombre_archivo = f"Carta_Invitacion_{i+1:02d}_{nombre_empresa_limpio}.docx"
                archivo_salida = os.path.join(directorio_cartas, nombre_archivo)
                
                datos_carta = self._preparar_datos_carta_empresa(contract_data, empresa, i)
                
                if self._generar_carta_individual(ruta_plantilla, archivo_salida, datos_carta):
                    cartas_generadas.append({
                        'archivo': archivo_salida,
                        'empresa': empresa.get('nombre', f'Empresa {i+1}'),
                        'numero': i + 1
                    })
            
            self._mostrar_resultado_cartas(cartas_generadas, contract_data, "Cartas de Invitación")
            if cartas_generadas:
                self._abrir_cartas_generadas(cartas_generadas)
                
                # 🆕 NOTIFICAR AL CRONOGRAMA QUE SE GENERARON LAS CARTAS DE INVITACIÓN
                self._actualizar_fase_en_generacion('cartas_invitacion')
            
        except Exception as e:
            self._mostrar_error(f"Error generando cartas de invitación: {str(e)}")

    def generar_acta_adjudicacion(self):
        """Generar acta de adjudicación"""
        self._generar_documento_simple('generar_acta_adjudicacion', 'Acta_Adjudicacion', "Acta de Adjudicación")
    
    def _limpiar_nombre_archivo(self, nombre):
        """Limpiar nombre para archivo"""
        try:
            # Usar la función de helpers si existe
            from helpers_py import limpiar_nombre_archivo
            return limpiar_nombre_archivo(nombre)
        except ImportError:
            # Implementación simple si no existe helpers
            import re
            nombre_limpio = re.sub(r'[<>:"/\\|?*]', '_', nombre)
            return nombre_limpio.strip('_')
    
    def generar_acta_liquidacion(self):
        """Generar acta de liquidación"""
        self._generar_documento_simple('generar_acta_liquidacion', 'Acta_Liquidacion', "Acta de Liquidación")

    def generar_acta_replanteo(self):
        """Generar acta de replanteo"""
        self._generar_documento_simple('generar_acta_replanteo', 'Acta_Replanteo', "Acta de Replanteo")

    def generar_acta_recepcion(self):
        """Generar acta de recepción"""
        self._generar_documento_simple('generar_acta_recepcion', 'Acta_Recepcion', "Acta de Recepción")

    def generar_nombramiento_director(self):
        """Generar nombramiento de director"""
        self._generar_documento_simple('generar_nombramiento_director', 'Nombramiento_Director', "Nombramiento Director")

    def generar_contrato(self):
        """Generar contrato CON selección dinámica de plantilla"""
        try:
            
            # 🆕 NUEVO: Verificación de plantilla dinámica
            nombre_plantilla = self._obtener_nombre_plantilla_dinamico('generar_contrato')
            if not self.verificar_plantilla_disponible(nombre_plantilla):
                # Fallback a plantilla legacy
                nombre_plantilla_legacy = 'plantilla_contrato.docx'
                if not self.verificar_plantilla_disponible(nombre_plantilla_legacy):
                    self._mostrar_error(f"No se encontraron plantillas: {nombre_plantilla} ni {nombre_plantilla_legacy}")
                    return
            
            # Continuar con la generación normal
            self._generar_documento_simple('generar_contrato', 'Contrato', "Contrato")
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error en generar_contrato: {e}")
            import traceback
            logger.exception("Error completo:")
            self._mostrar_error(f"Error inesperado generando Contrato: {str(e)}")

    # =================== MÉTODOS AUXILIARES PARA GENERACIÓN ===================

    def _generar_documento_simple(self, tipo_funcion: str, nombre_archivo: str, titulo: str):
        """Generar documento simple CON TODAS LAS VERIFICACIONES Y TRACKING"""
        documento_id = None
        try:
            if not self._validar_contrato_seleccionado():
                logger.warning(f"[ControladorDocumentos] ❌ Validación falló")
                return False
            
            contract_data = self._obtener_datos_contrato_actual()
            if not contract_data:
                return self._mostrar_error("No se pudieron obtener los datos del contrato")
            
            # 🆕 TRACKING: Iniciar seguimiento del documento
            tipo_documento = self._mapear_tipo_documento(tipo_funcion)
            plantilla = self._obtener_nombre_plantilla_dinamico(tipo_funcion)
            documento_id = self._iniciar_tracking_documento(tipo_documento, titulo, plantilla)
            
            # 🎯 ACTUALIZAR FASE: Marcar fase como generada ANTES de generar
            self._actualizar_fase_en_generacion(tipo_documento)
            
            # Generar documento
            exito, ruta_generada = self._generar_documento_con_sustitucion_tracked(
                tipo_funcion, contract_data, nombre_archivo
            )
            
            if exito and ruta_generada:
                # 🆕 TRACKING: Completar seguimiento
                self._completar_tracking_documento(
                    documento_id, 
                    ruta_generada, 
                    f"Generado exitosamente usando {plantilla or 'plantilla por defecto'}"
                )
                
                QMessageBox.information(
                    self.main_window, titulo, 
                    f"✅ {titulo} generado correctamente"
                )
                
                # Mostrar diálogo de conversión a PDF
                self._mostrar_dialogo_conversion_pdf(titulo, ruta_generada)
                
                return True
            else:
                # 🆕 TRACKING: Registrar error
                self._error_tracking_documento(documento_id, "Error en la generación del documento")
                self._mostrar_error(f"Error generando {titulo}")
                return False
                
        except Exception as e:
            # 🆕 TRACKING: Registrar excepción
            if documento_id:
                self._error_tracking_documento(documento_id, f"Excepción: {str(e)}")
            
            logger.error(f"[ControladorDocumentos] ❌ Excepción: {e}")
            import traceback
            logger.exception("Error completo:")
            return False

    def _generar_documento_con_sustitucion(self, tipo_funcion: str, contract_data: Dict[str, Any], nombre_archivo: str) -> bool:
        """Método principal para generar documentos - CON SELECCIÓN DINÁMICA"""
        try:
            
            # Obtener plantilla dinámicamente según tipo de contrato
            nombre_plantilla = self._obtener_nombre_plantilla_dinamico(tipo_funcion)
            if not nombre_plantilla:
                logger.error(f"[ControladorDocumentos] ❌ No se encontró plantilla para: {tipo_funcion}")
                return False
            
            logger.info(f"[ControladorDocumentos] 📄 Usando plantilla: {nombre_plantilla}")
            
            ruta_plantilla = self._obtener_ruta_plantilla(nombre_plantilla)
            if not ruta_plantilla:
                logger.error(f"[ControladorDocumentos] ❌ No se encontró archivo de plantilla: {nombre_plantilla}")
                # 🆕 FALLBACK: Intentar con plantilla legacy
                nombre_plantilla_legacy = self.mapeo_plantillas.get(tipo_funcion)
                if nombre_plantilla_legacy:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Fallback a plantilla legacy: {nombre_plantilla_legacy}")
                    ruta_plantilla = self._obtener_ruta_plantilla(nombre_plantilla_legacy)
                    
                if not ruta_plantilla:
                    logger.error(f"[ControladorDocumentos] ❌ No se encontró archivo de plantilla ni legacy")
                    return False
            
            # 🆕 USAR GESTOR UNIFICADO PARA OBTENER CARPETA
            carpeta_contrato = self._obtener_carpeta_con_gestor_unificado(contract_data)
            if not carpeta_contrato:
                logger.error(f"[ControladorDocumentos] ❌ No se pudo obtener carpeta del contrato")
                return False
            
            # 🆕 DETERMINAR SUBCARPETA SEGÚN TIPO DE DOCUMENTO
            subcarpeta_destino = self._determinar_subcarpeta_por_tipo_documento(tipo_funcion)
            
            # Construir ruta completa
            nombre_archivo_limpio = self._limpiar_nombre_archivo(f"{nombre_archivo}")
            archivo_salida = os.path.join(carpeta_contrato, subcarpeta_destino, f"{nombre_archivo_limpio}.docx")
            
            
            # Crear directorio de destino
            directorio_documentos = os.path.dirname(archivo_salida)
            os.makedirs(directorio_documentos, exist_ok=True)
            
            # Generar documento
            exito = self._sustituir_variables_en_documento(ruta_plantilla, archivo_salida, contract_data)
            
            if exito and os.path.exists(archivo_salida):
                self._abrir_documento(archivo_salida)
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error generando documento: {e}")
            return False

    def _sustituir_variables_en_documento(self, ruta_plantilla: str, archivo_salida: str, contract_data: Dict[str, Any]) -> bool:
        """Sustituir variables en documento Word - CON DETECCIÓN DE CAMPOS VACÍOS"""
        try:
            
            # Verificar plantilla
            if not os.path.exists(ruta_plantilla):
                logger.error(f"[ControladorDocumentos] ❌ Plantilla no existe: {ruta_plantilla}")
                return False
            
            # Preparar datos para sustitución
            datos_completos = self._preparar_datos_para_sustitucion(contract_data)
            
            
            # 🆕 NUEVO: Detectar variables en plantilla y verificar campos vacíos
            variables_en_plantilla = self._detectar_variables_en_plantilla(ruta_plantilla)
            campos_vacios = self._verificar_campos_vacios(variables_en_plantilla, datos_completos)
            
            if campos_vacios:
                if not self._mostrar_popup_campos_vacios(campos_vacios, os.path.basename(archivo_salida)):
                    logger.warning(f"[ControladorDocumentos] ⚠️ Generación cancelada por campos vacíos")
                    return False
            
            # Abrir documento
            doc = Document(ruta_plantilla)
            
            # Procesar párrafos principales
            variables_encontradas = set()
            
            for i, paragraph in enumerate(doc.paragraphs):
                try:
                    vars_parrafo = self._procesar_paragraph_con_variables(paragraph, datos_completos)
                    variables_encontradas.update(vars_parrafo)
                except Exception as e:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Error en párrafo {i}: {e}")
                    continue
            
            # Procesar tablas
            for i, table in enumerate(doc.tables):
                try:
                    for j, row in enumerate(table.rows):
                        for k, cell in enumerate(row.cells):
                            for l, paragraph in enumerate(cell.paragraphs):
                                vars_celda = self._procesar_paragraph_con_variables(paragraph, datos_completos)
                                variables_encontradas.update(vars_celda)
                except Exception as e:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Error en tabla {i}: {e}")
                    continue
            
            # Procesar headers y footers
            try:
                for section in doc.sections:
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            vars_header = self._procesar_paragraph_con_variables(paragraph, datos_completos)
                            variables_encontradas.update(vars_header)
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            vars_footer = self._procesar_paragraph_con_variables(paragraph, datos_completos)
                            variables_encontradas.update(vars_footer)
            except Exception as e:
                logger.warning(f"[ControladorDocumentos] ⚠️ Error en headers/footers: {e}")
            
            # 🆕 NUEVO: Procesar marcadores especiales de tabla (como antes)
            try:
                empresas_lista = self._obtener_empresas_lista(contract_data)
                if empresas_lista:
                    self._sustituir_marcadores_tabla(doc, empresas_lista)
            except Exception as e:
                logger.warning(f"[ControladorDocumentos] ⚠️ Error en tablas especiales: {e}")
            
            # Crear directorio de salida
            directorio_salida = os.path.dirname(archivo_salida)
            os.makedirs(directorio_salida, exist_ok=True)
            
            # Guardar documento
            doc.save(archivo_salida)
            
            
            if variables_encontradas:
                pass  # Variables found successfully
            
            return True
            
        except Exception as e:
            logger.critical(f"[ControladorDocumentos] ❌ Error crítico procesando documento: {e}")
            import traceback
            logger.exception("Error completo:")
            
            error_msg = f"Error procesando documento:\n{str(e)}\n\nVerifica que:\n"
            error_msg += "• La plantilla no esté abierta en Word\n"
            error_msg += "• Tengas permisos de escritura\n"
            error_msg += "• El archivo no esté corrupto"
            
            self._mostrar_error(error_msg)
            return False
    def _obtener_empresas_lista(self, contract_data: Dict[str, Any]) -> List[Dict]:
        """Obtener lista de empresas del contrato"""
        try:
           # logger.info(f"[DEBUG] CONTRACT_DATA EMPRESAS: {contract_data.get('empresas', {})}")
            
            empresas = contract_data.get('empresas', {})
           # logger.info(f"[DEBUG] Tipo de empresas: {type(empresas)}")
            if isinstance(empresas, dict):
                empresas_lista = empresas.get('empresa', [])
            elif isinstance(empresas, list):
                empresas_lista = empresas
            else:
                empresas_lista = []
                
           # logger.info(f"[DEBUG] EMPRESAS LISTA FINAL ({len(empresas_lista)}): {empresas_lista}")
            return empresas_lista
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error obteniendo lista de empresas: {e}")
            return []
    def _detectar_variables_en_plantilla(self, ruta_plantilla: str) -> set:
        """Detectar qué variables están presentes en la plantilla"""
        try:
            from docx import Document
            doc = Document(ruta_plantilla)
            variables_encontradas = set()
            
            import re
            patron = r'@(\w+)@'
            
            # Buscar en párrafos
            for paragraph in doc.paragraphs:
                variables_encontradas.update(re.findall(patron, paragraph.text))
            
            # Buscar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            variables_encontradas.update(re.findall(patron, paragraph.text))
            
            # Buscar en headers y footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        variables_encontradas.update(re.findall(patron, paragraph.text))
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        variables_encontradas.update(re.findall(patron, paragraph.text))
            
            return variables_encontradas
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error detectando variables: {e}")
            return set()

    def _verificar_campos_vacios(self, variables_plantilla: set, datos_disponibles: dict) -> list:
        """Verificar qué campos están vacíos o no existen - VERSION MEJORADA"""
        campos_vacios = []
        
        #logger.debug(f"[DEBUG] VERIFICANDO CAMPOS VACÍOS")
        #logger.info(f"[DEBUG] Variables en plantilla: {len(variables_plantilla)}")
        #logger.info(f"[DEBUG] Datos disponibles: {len(datos_disponibles)}")
        
        for variable in variables_plantilla:
            # Ignorar marcadores especiales de tabla
            if variable.startswith('tabla-') or variable == 'tablaAnualidades':
                continue
            
            #logger.info(f"[DEBUG] 🔎 Verificando variable: '{variable}'")
            
            # Verificar si el campo no existe
            if variable not in datos_disponibles:
                campos_vacios.append(f"{variable} (no existe)")
                #logger.error(f"[DEBUG] '{variable}' NO EXISTE en datos")
            else:
                # Verificar si está vacío
                valor = datos_disponibles[variable]
                
                # Considerar vacío si es None, string vacío, o valores placeholder
                valores_vacios = [None, "", " ", "---------------------", "--------------------", "0", "0.0", "0.00"]
                
                if valor in valores_vacios or str(valor).strip() == "":
                    campos_vacios.append(f"{variable} (vacío: '{valor}')")
                    #logger.warning(f"[DEBUG] '{variable}' ESTÁ VACÍO: '{valor}'")
                else:
                    #logger.info(f"[DEBUG] '{variable}' TIENE VALOR: '{str(valor)[:30]}{'...' if len(str(valor)) > 30 else ''}'")
                    pass
        #logger.info(f"[DEBUG] RESULTADO: {len(campos_vacios)} campos vacíos de {len(variables_plantilla)} variables")
        

        
        return campos_vacios
    def _mostrar_popup_campos_vacios(self, campos_vacios: list, nombre_documento: str) -> bool:
        """Mostrar popup con campos vacíos y preguntar si continuar"""
        try:
            from PyQt5.QtWidgets import QMessageBox
            
            # Preparar mensaje
            campos_texto = "\n".join([f"• {campo}" for campo in campos_vacios[:15]])
            if len(campos_vacios) > 15:
                campos_texto += f"\n• ... y {len(campos_vacios) - 15} campos más"
            
            msg = QMessageBox(self.main_window)
            msg.setIcon(QMessageBox.Warning)
            msg.setWindowTitle("⚠️ Campos Vacíos Detectados")
            msg.setText(f"<b>El documento '{nombre_documento}' tiene {len(campos_vacios)} campos vacíos</b>")
            msg.setInformativeText(
                f"<b>Campos vacíos o no encontrados:</b><br><br>"
                f"<font color='#666666'>{campos_texto.replace(chr(10), '<br>')}</font><br><br>"
                f"<b>¿Deseas continuar generando el documento?</b><br>"
                f"• <b>Sí:</b> Generar documento con campos vacíos<br>"
                f"• <b>No:</b> Cancelar y completar datos primero"
            )
            
            msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg.setDefaultButton(QMessageBox.No)
            
            # Personalizar botones
            yes_button = msg.button(QMessageBox.Yes)
            no_button = msg.button(QMessageBox.No)
            yes_button.setText("✅ Continuar")
            no_button.setText("❌ Cancelar")
            
            resultado = msg.exec_()
            
            if resultado == QMessageBox.Yes:
                logger.warning(f"[ControladorDocumentos] ⚠️ Usuario eligió continuar con {len(campos_vacios)} campos vacíos")
                return True
            else:
                logger.info(f"[ControladorDocumentos] ❌ Usuario canceló por campos vacíos")
                return False
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error mostrando popup: {e}")
            return True  # En caso de error, continuar

    
    def _preparar_datos_para_sustitucion(self, contract_data, usar_firmantes_globales=True):
        """Preparar datos para sustitución en documentos Word"""
        try:
            # Campos de firmantes que se manejan globalmente
            campos_firmantes = {
                'firmanteConforme', 'cargoConforme', 'firmantePropone', 'cargoPropone',
                'firmanteAprueba', 'cargoAprueba', 'representanteFirma', 'cargoResponsable',
                'representanteAdif', 'cargoResponsable1', 'firmanteContratacion', 'cargoContratacion',
                'firmanteResponsable', 'directorFacultativo', 'nombreAsistenteAdif'
            }
            
            # Filtrar datos del contrato excluyendo firmantes individuales
            datos_filtrados = {}
            for campo, valor in contract_data.items():
                if campo not in campos_firmantes:
                    datos_filtrados[campo] = valor
                #else:
                    #logger.info(f"[DEBUG] 🚫 Firmante individual IGNORADO: {campo} = '{valor}' (será reemplazado por global)")
            
            #logger.info(f"[DEBUG] PREPARANDO DATOS PARA SUSTITUCIÓN")
            #logger.info(f"[DEBUG] 📋 Datos filtrados (sin firmantes individuales): {len(datos_filtrados)} campos")
            
            # Inicializar datos finales
            datos_finales = datos_filtrados.copy()
            
            # Cargar y agregar firmantes globales
            if usar_firmantes_globales:
                try:
                    # Buscar archivo JSON
                    current_dir = os.path.dirname(__file__)
                    parent_dir = os.path.dirname(current_dir)
                    # Intentar primero en carpeta basedatos
                    json_path = os.path.join(parent_dir, "basedatos", "BaseDatos.json")
                    
                    # Si no existe, intentar en raíz como fallback
                    if not os.path.exists(json_path):
                        json_path = os.path.join(parent_dir, "BaseDatos.json")
                    
                    if os.path.exists(json_path):
                        with open(json_path, "r", encoding="utf-8") as f:
                            data = json.load(f)
                        
                        firmantes_globales = data.get('firmantes', {})
                        #logger.info(f"[DEBUG] 👥 Firmantes globales encontrados: {len(firmantes_globales)}")
                        
                        # Agregar cada firmante global
                        for campo_firmante, valor_firmante in firmantes_globales.items():
                            if valor_firmante and valor_firmante.strip():
                                datos_finales[campo_firmante] = valor_firmante
                                #logger.info(f"[DEBUG] Firmante GLOBAL agregado: {campo_firmante} = '{valor_firmante}'")
                            #else:
                                #logger.info(f"[DEBUG] 🔳 Firmante GLOBAL vacío: {campo_firmante} = '{valor_firmante}'")
                    
                except Exception as e:
                    logger.error(f"[DEBUG] ❌ Error cargando firmantes globales: {e}")
            
            # 🆕 PROCESAR CAMPOS ENTEROS CON CARACTERES INVISIBLES (CORREGIDO)
            campos_enteros = ['plazoEjecucion', 'numEmpresasPresentadas', 'numEmpresasSolicitadas']
            for campo in campos_enteros:
                if campo in datos_finales:
                    try:
                        valor_original = datos_finales[campo]
                        #logger.info(f"[DEBUG] 🔧 {campo} ORIGINAL: '{valor_original}' (tipo: {type(valor_original)})")
                        
                        # Convertir a entero puro
                        if valor_original is not None and str(valor_original).strip():
                            valor_limpio = str(valor_original).replace(',', '.').strip()
                            entero = int(float(valor_limpio))
                            
                            # SOLUCIÓN CORREGIDA: Siempre agregar caracteres invisibles
                            numero_str = str(entero)
                            if len(numero_str) == 1:
                                # Para números de un dígito, agregar caracteres invisibles al final
                                resultado = numero_str + "\u200B\u2060\u200C"
                            else:
                                # Para números de múltiples dígitos, intercalar entre dígitos
                                resultado = ""
                                for i, digito in enumerate(numero_str):
                                    resultado += digito
                                    if i < len(numero_str) - 1:
                                        resultado += "\u200B\u2060\u200C"
                            
                            datos_finales[campo] = resultado
                            if campo == 'plazoEjecucion':
                                datos_finales['plazoEjecucionTexto'] = resultado
                            
                            #logger.info(f"[DEBUG] 🔧 {campo} DEFINITIVO: '{resultado}' (con caracteres invisibles)")
                            #logger.info(f"[DEBUG] 🔧 Longitud: {len(resultado)} caracteres")
                        else:
                            datos_finales[campo] = "0\u200B"  # Incluso el 0 con caracteres invisibles
                            if campo == 'plazoEjecucion':
                                datos_finales['plazoEjecucionTexto'] = "0\u200B"
                                
                    except Exception as e:
                        logger.error(f"[DEBUG] ❌ Error procesando {campo}: {e}")
                        datos_finales[campo] = "0\u200B"
                        if campo == 'plazoEjecucion':
                            datos_finales['plazoEjecucionTexto'] = "0\u200B"
            
            # Procesar campos monetarios normalmente
            campos_numericos_otros = ['basePresupuesto', 'totalPresupuestoBase', 'ivaPresupuestoBase', 
                                    'precioAdjudicacion', 'precioAdjudicacionIva', 'precioAdjudicacionTotal']
            for campo in campos_numericos_otros:
                if campo in datos_finales:
                    try:
                        valor = datos_finales[campo]
                        if valor is not None and str(valor).strip():
                            decimal = float(str(valor).replace(',', '.'))
                            datos_finales[campo] = f"{decimal:.2f}"
                            #logger.info(f"[DEBUG] 💰 {campo} formateado como: '{datos_finales[campo]}'")
                    except (ValueError, TypeError):
                        datos_finales[campo] = "0.00"
            
            #logger.info(f"[DEBUG] RESUMEN FINAL:")
            #logger.info(f"[DEBUG] Total campos preparados: {len(datos_finales)}")
            #logger.info(f"[DEBUG] 🔧 Enteros con caracteres invisibles para forzar texto")
            
            return datos_finales
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error preparando datos: {e}")
            import traceback
            logger.exception("Error completo:")
            return contract_data  # Devolver datos originales en caso de error

    # =================== MÉTODOS DE CARPETAS Y ARCHIVOS ===================

    def _crear_o_obtener_carpeta_contrato(self, contract_data: Dict[str, Any]) -> Optional[str]:
        """Crear o obtener la carpeta del contrato"""
        try:
            # Usar controlador de archivos de la ventana principal
            if (self.main_window and 
                hasattr(self.main_window, 'controlador_archivos') and 
                self.main_window.controlador_archivos):
                
                carpeta_path = self.main_window.controlador_archivos.crear_carpeta_contrato_actual()
                if carpeta_path:
                    return carpeta_path
            
            # Usar contract_manager
            if (self.main_window and 
                hasattr(self.main_window, 'contract_manager') and 
                self.main_window.contract_manager):
                
                carpeta_path = self.main_window.contract_manager.crear_carpeta_para_contrato_actual()
                if carpeta_path:
                    return carpeta_path
            
            # Fallback: crear carpeta manualmente
            return self._crear_carpeta_manual_simple(contract_data)
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error creando carpeta de contrato: {e}")
            return self._crear_carpeta_manual_simple(contract_data)

    

    def _obtener_ruta_plantilla(self, nombre_plantilla: str) -> Optional[str]:
        """Obtiene la ruta completa de una plantilla usando el controlador de rutas"""
        try:
            from .controlador_routes import rutas
            return rutas.get_ruta_plantilla(nombre_plantilla)
        except ImportError:
            # Fallback al método original si no está disponible el controlador
            ubicaciones = [
                os.path.join(self.directorio_plantillas, nombre_plantilla),
                os.path.join("plantillas", nombre_plantilla),
                os.path.join("templates", nombre_plantilla),
                nombre_plantilla
            ]
            
            for ubicacion in ubicaciones:
                if os.path.exists(ubicacion):
                    return ubicacion
            
            return None

    def _limpiar_nombre_para_archivo(self, nombre):
        """Limpiar nombre de empresa para usar como nombre de archivo"""
        try:
            nombre_limpio = re.sub(r'[<>:"/\\|?*]', '_', nombre)
            nombre_limpio = re.sub(r'\s+', '_', nombre_limpio)
            nombre_limpio = nombre_limpio.strip('_')
            
            if len(nombre_limpio) > 50:
                nombre_limpio = nombre_limpio[:47] + "..."
            
            return nombre_limpio if nombre_limpio else "empresa"
        except Exception:
            return "empresa"

    def _abrir_documento(self, ruta_archivo: str):
        """Abrir documento generado"""
        try:
            if os.path.exists(ruta_archivo):
                if sys.platform.startswith('win'):
                    os.startfile(ruta_archivo)
                elif sys.platform.startswith('darwin'):
                    subprocess.run(['open', ruta_archivo])
                else:
                    subprocess.run(['xdg-open', ruta_archivo])
        except Exception:
            pass

    def _abrir_cartas_generadas(self, cartas_generadas):
        """Abrir todas las cartas generadas"""
        try:
            
            for carta in cartas_generadas:
                self._abrir_documento(carta['archivo'])
        except Exception:
            pass

    # =================== MÉTODOS DE VALIDACIÓN Y MENSAJES ===================

    def _validar_contrato_seleccionado(self) -> bool:
        """Validar que hay un contrato seleccionado"""
        if not self.main_window:
            return False
        
        if not hasattr(self.main_window, 'contract_manager') or not self.main_window.contract_manager:
            self._mostrar_error("Gestor de contratos no disponible")
            return False
        
        if not self.main_window.contract_manager.get_current_contract():
            QMessageBox.warning(
                self.main_window, "Advertencia", 
                "⚠️ Debes seleccionar un contrato primero"
            )
            return False
        
        return True

    def _obtener_datos_contrato_actual(self) -> Optional[Dict[str, Any]]:
        """Obtener datos del contrato actual - LEYENDO JSON DIRECTAMENTE"""
        
        if not self.main_window:
            logger.error(f"[ControladorDocumentos] ❌ main_window es None")
            return None
        
        if not hasattr(self.main_window, 'contract_manager'):
            logger.error(f"[ControladorDocumentos] ❌ No tiene contract_manager")
            return None
        
        if not self.main_window.contract_manager:
            logger.error(f"[ControladorDocumentos] ❌ contract_manager es None")
            return None
        
        # Obtener nombre del contrato actual
        nombre_contrato = self.main_window.contract_manager.get_current_contract()
        if not nombre_contrato:
            logger.warning(f"[ControladorDocumentos] ❌ No hay contrato seleccionado")
            return None
        
        # 🆕 NUEVO: Guardar nombre del contrato para tracking
        self.contract_name = nombre_contrato
        
        
        # 🆕 LEER DIRECTAMENTE DEL JSON EN LUGAR DE CACHE
        if (hasattr(self.main_window, 'controlador_json') and 
            self.main_window.controlador_json and
            self.main_window.controlador_json.gestor):
            
            
            # Recargar datos del JSON desde disco
            self.main_window.controlador_json.gestor.recargar_datos()
            
            # Leer contrato completo del JSON
            data_json = self.main_window.controlador_json.leer_contrato_completo(nombre_contrato)
            
            if data_json:
                
                # Debug de campos críticos
                campos_debug = ['nombreObra', 'plazoEjecucion', 'representanteAdif', 'basePresupuesto']
                for campo in campos_debug:
                    valor = data_json.get(campo, 'N/A')
                
                return data_json
            else:
                logger.error(f"[ControladorDocumentos] ❌ No se encontraron datos en JSON para: {nombre_contrato}")
        
        # Fallback al método original solo si falla el JSON
        data = self.main_window.contract_manager.get_current_contract_data()
        
        if data:
            return data
        else:
            logger.error(f"[ControladorDocumentos] ❌ get_current_contract_data() retornó None")
            return None
    def _mostrar_error(self, mensaje: str):
        """Mostrar mensaje de error"""
        if self.main_window:
            QMessageBox.critical(self.main_window, "❌ Error", mensaje)
            
    def _mostrar_dialogo_conversion_pdf(self, titulo: str, ruta_docx: str):
        """Conversión automática a PDF sin diálogo"""
        try:
            if ruta_docx and ruta_docx.endswith('.docx') and os.path.exists(ruta_docx):
                logger.info(f"[ControladorDocumentos] 📄 Generando PDF automáticamente para: {ruta_docx}")
                
                # Importar función de conversión
                from .controlador_pdf_unificado import convertir_docx_a_pdf_simple
                
                # Convertir directamente
                if convertir_docx_a_pdf_simple(ruta_docx):
                    pdf_path = ruta_docx.replace('.docx', '.pdf')
                    
                    # Abrir automáticamente el PDF
                    import subprocess
                    try:
                        logger.info(f"[ControladorDocumentos] 📄 Abriendo PDF automáticamente: {pdf_path}")
                        subprocess.run([pdf_path], shell=True, check=True)
                    except Exception as e:
                        logger.warning(f"[ControladorDocumentos] ⚠️ Error abriendo PDF: {e}")
                        # Si falla abrir el PDF, abrir la carpeta
                        carpeta = os.path.dirname(ruta_docx)
                        subprocess.run(f'explorer "{carpeta}"', shell=True)
                else:
                    logger.error(f"[ControladorDocumentos] ❌ Error generando PDF para: {ruta_docx}")
            else:
                logger.warning(f"[ControladorDocumentos] ⚠️ No se puede mostrar diálogo PDF - Archivo no válido: {ruta_docx}")
        except ImportError as e:
            logger.error(f"[ControladorDocumentos] ❌ Error de importación en diálogo PDF: {e}")
            if self.main_window:
                from PyQt5.QtWidgets import QMessageBox
                QMessageBox.warning(self.main_window, "Advertencia", 
                                  "El módulo de conversión PDF no está disponible")
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error inesperado en diálogo PDF: {e}")
            import traceback
            logger.exception("Error completo:")

    def _mostrar_resultado_cartas(self, cartas_generadas, contract_data, tipo_carta):
        """Mostrar resultado de la generación de cartas"""
        try:
            nombre_contrato = contract_data.get('nombreObra', 'Sin nombre')
            total_generadas = len(cartas_generadas)
            
            if total_generadas > 0:
                mensaje = f"✅ {tipo_carta.upper()} GENERADAS\n\n"
                mensaje += f"📂 Contrato: {nombre_contrato}\n"
                mensaje += f"📄 Cartas generadas: {total_generadas}\n\n"
                mensaje += "📋 DETALLE:\n"
                
                for carta in cartas_generadas[:5]:
                    mensaje += f"• {carta['numero']:02d}. {carta['empresa']}\n"
                
                if total_generadas > 5:
                    mensaje += f"• ... y {total_generadas - 5} cartas más\n"
                
                QMessageBox.information(self.main_window, tipo_carta, mensaje)
            else:
                mensaje_error = f"❌ NO SE GENERARON {tipo_carta.upper()}\n\n"
                mensaje_error += "No se encontraron empresas válidas para generar cartas."
                QMessageBox.critical(self.main_window, "Error", mensaje_error)
                
        except Exception:
            pass


    # =================== FUNCIONES DE UTILIDAD ===================

    def formatear_numero_para_documento(numero: float) -> str:
        """Formatea número para documentos en formato español"""
        if numero >= 1000:
            return f"{numero:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return f"{numero:.2f}".replace(".", ",")


    def validar_plantilla_existe(nombre_plantilla: str) -> bool:
        """Valida que una plantilla existe"""
        ubicaciones = [
            resource_path(f"templates/{nombre_plantilla}"),
            resource_path(f"plantillas/{nombre_plantilla}"),
            os.path.join("plantillas", nombre_plantilla),
            nombre_plantilla
        ]
        
        return 

    def _sustituir_marcadores_tabla(self, doc: Document, empresas_lista: List[Dict]):
        """Sustituir marcadores @tabla-ofertas@ en TODO el documento"""
        try:
            logger.debug(f"[DEBUG] SUSTITUYENDO MARCADORES - EMPRESAS: {empresas_lista}")
            
            paragrafos_a_procesar = []
            
            # 1. BUSCAR EN PÁRRAFOS PRINCIPALES
            for i, paragraph in enumerate(doc.paragraphs):
                texto = paragraph.text
                if '@tabla-ofertas@' in texto:
                    logger.debug(f"[DEBUG] ENCONTRADO @tabla-ofertas@ EN PÁRRAFO PRINCIPAL {i}")
                    paragrafos_a_procesar.append((paragraph, 'ofertas'))
                elif '@tabla-empresas@' in texto:
                    paragrafos_a_procesar.append((paragraph, 'empresas'))
                elif '@tabla-clasificacion@' in texto:
                    paragrafos_a_procesar.append((paragraph, 'clasificacion'))
            
            # 2. BUSCAR EN TABLAS
            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        for l, paragraph in enumerate(cell.paragraphs):
                            texto = paragraph.text
                            if '@tabla-ofertas@' in texto:
                                logger.debug(f"[DEBUG] ENCONTRADO @tabla-ofertas@ EN TABLA {i}, FILA {j}, CELDA {k}")
                                paragrafos_a_procesar.append((paragraph, 'ofertas'))
                            elif '@tabla-empresas@' in texto:
                                paragrafos_a_procesar.append((paragraph, 'empresas'))
                            elif '@tabla-clasificacion@' in texto:
                                paragrafos_a_procesar.append((paragraph, 'clasificacion'))
                            elif '@tablaAnualidades@' in texto:
                                logger.debug(f"[DEBUG] ENCONTRADO @tablaAnualidades@ EN TABLA {i}, FILA {j}, CELDA {k}")
                                paragrafos_a_procesar.append((paragraph, 'anualidades'))
            
            # 3. BUSCAR EN HEADERS Y FOOTERS
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        texto = paragraph.text
                        if '@tabla-ofertas@' in texto:
                            logger.debug(f"[DEBUG] ENCONTRADO @tabla-ofertas@ EN HEADER")
                            paragrafos_a_procesar.append((paragraph, 'ofertas'))
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        texto = paragraph.text
                        if '@tabla-ofertas@' in texto:
                            logger.debug(f"[DEBUG] ENCONTRADO @tabla-ofertas@ EN FOOTER")
                            paragrafos_a_procesar.append((paragraph, 'ofertas'))
            
            logger.debug(f"[DEBUG] PÁRRAFOS A PROCESAR: {len(paragrafos_a_procesar)}")
            
            # 4. PROCESAR TODOS LOS MARCADORES ENCONTRADOS
            for paragraph, tipo_tabla in paragrafos_a_procesar:
                if tipo_tabla == 'ofertas':
                    logger.debug(f"[DEBUG] LLAMANDO _insertar_tabla_ofertas")
                    self._insertar_tabla_ofertas(doc, paragraph, empresas_lista)
                elif tipo_tabla == 'empresas':
                    self._insertar_tabla_empresas(doc, paragraph, empresas_lista)
                elif tipo_tabla == 'clasificacion':
                    self._insertar_tabla_clasificacion(doc, paragraph, empresas_lista)
                elif tipo_tabla == 'anualidades':
                    logger.debug(f"[DEBUG] LLAMANDO _insertar_tabla_anualidades")
                    self._insertar_tabla_anualidades(doc, paragraph)
            
        except Exception as e:
            logger.error(f"[DEBUG] ERROR EN MARCADORES: {e}")
            import traceback
            logger.exception("Error completo:")
    def obtener_empresas_para_docx(self, contract_data):
        """NUEVA: Obtener empresas en formato para documentos DOCX"""
        try:
            # ✅ OBTENER EMPRESAS UNIFICADAS
            empresas_lista = contract_data.get('empresas', [])
            
            # Migrar estructura antigua si es necesaria
            if isinstance(empresas_lista, dict) and 'empresa' in empresas_lista:
                empresas_lista = self._migrar_estructura_para_docx(empresas_lista, contract_data)

            empresas_para_docx = []
            
            for empresa in empresas_lista:
                if isinstance(empresa, dict):
                    empresa_docx = {
                        'nombre': empresa.get('nombre', ''),
                        'nif': empresa.get('nif', ''),
                        'email': empresa.get('email', ''),
                        'contacto': empresa.get('contacto', ''),
                        'ofertas': empresa.get('ofertas', ''),
                        # Campos adicionales para documentos si es necesario
                        'ofertas_formateada': self._formatear_oferta_euros(empresa.get('ofertas', '')),
                    }
                    empresas_para_docx.append(empresa_docx)

            return empresas_para_docx

        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error preparando empresas para DOCX: {e}")
            return []

    def _formatear_oferta_euros(self, oferta_str):
        """NUEVA: Formatear oferta para documentos"""
        try:
            if not oferta_str or oferta_str.strip() == '':
                return "No ofertado"
            
            # Convertir a número y formatear
            valor = float(oferta_str.replace(',', '.'))
            return f"{valor:,.2f} €".replace(',', '.')
            
        except (ValueError, TypeError):
            return oferta_str

    def _migrar_estructura_para_docx(self, empresas_dict, contract_data):
        """NUEVA: Migrar estructura antigua para DOCX"""
        try:
            empresas_antiguas = empresas_dict.get('empresa', [])
            ofertas_antiguas = contract_data.get('ofertas', [])
            
            empresas_unificadas = []
            
            for i, empresa in enumerate(empresas_antiguas):
                empresa_unificada = {
                    'nombre': empresa.get('nombre', empresa.get('empresa', '')),
                    'nif': empresa.get('nif', empresa.get('cif', '')),
                    'email': empresa.get('email', ''),
                    'contacto': empresa.get('contacto', empresa.get('persona de contacto', '')),
                    'ofertas': ''
                }
                
                # Buscar oferta correspondiente
                if i < len(ofertas_antiguas):
                    oferta_antigua = ofertas_antiguas[i]
                    if isinstance(oferta_antigua, dict):
                        empresa_unificada['ofertas'] = oferta_antigua.get('oferta_(€)', '')
                
                empresas_unificadas.append(empresa_unificada)
            
            return empresas_unificadas
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error migrando para DOCX: {e}")
            return []

    def _insertar_tabla_empresas(self, doc: Document, paragraph, empresas_lista: List[Dict]):
        """Insertar tabla de empresas - VERSIÓN MEJORADA"""
        try:
            logger.debug(f"[DEBUG] 📊 INSERTANDO TABLA EMPRESAS")
            logger.debug(f"[DEBUG] Empresas recibidas: {empresas_lista}")
            
            if not empresas_lista:
                paragraph.text = paragraph.text.replace('@tabla-empresas@', '[No hay empresas registradas]')
                return

            # Crear tabla con 4 columnas: Nombre, NIF, Email, Contacto
            tabla = doc.add_table(rows=len(empresas_lista) + 1, cols=4)
            self._configurar_tabla_basica(tabla)

            # Encabezados
            encabezados = ["NOMBRE DE LA EMPRESA", "NIF DE LA EMPRESA", "EMAIL", "PERSONA DE CONTACTO"]
            fila_enc = tabla.rows[0]
            
            for j, encabezado in enumerate(encabezados):
                celda = fila_enc.cells[j]
                celda.text = ""
                run = celda.paragraphs[0].add_run(encabezado)
                run.font.size = Pt(10)
                run.bold = True

            # Datos de empresas
            for i, empresa in enumerate(empresas_lista):
                logger.debug(f"[DEBUG] Procesando empresa {i+1}: {empresa}")
                
                # Obtener datos con valores por defecto
                nombre = empresa.get('nombre') or empresa.get('empresa') or f'Empresa {i+1}'
                nif = empresa.get('nif') or empresa.get('cif') or ''
                email = empresa.get('email') or ''
                contacto = empresa.get('contacto') or empresa.get('persona de contacto') or ''
                
                datos = [nombre, nif, email, contacto]
                logger.debug(f"[DEBUG] Datos extraídos: {datos}")

                fila = tabla.rows[i + 1]
                for j, dato in enumerate(datos):
                    celda = fila.cells[j]
                    celda.text = ""
                    run = celda.paragraphs[0].add_run(str(dato))
                    run.font.size = Pt(10)

            # Insertar tabla y eliminar marcador
            self._insertar_tabla_despues_de_parrafo(paragraph, tabla)
            paragraph.text = paragraph.text.replace('@tabla-empresas@', '')

        except Exception as e:
            logger.error(f"[DEBUG] ❌ ERROR insertando tabla: {e}")
            paragraph.text = paragraph.text.replace('@tabla-empresas@', '[Error insertando tabla de empresas]')

    def _insertar_tabla_ofertas(self, doc: Document, paragraph, empresas_lista: List[Dict]):
        """Crear tabla de ofertas simple leyendo empresas del JSON"""
        try:
            logger.debug(f"[DEBUG] 📊 CREANDO TABLA DE OFERTAS")
            logger.debug(f"[DEBUG] Empresas recibidas: {len(empresas_lista)}")
            
            # 1. PROCESAR DATOS DE EMPRESAS
            datos_tabla = []
            for i, empresa in enumerate(empresas_lista):
                nombre = empresa.get('nombre', f'Empresa {i+1}')
                ofertas_valor = empresa.get('ofertas', '')
                
                # Determinar si presenta oferta
                presenta_oferta = "No"
                precio_numerico = 0
                
                if ofertas_valor and str(ofertas_valor).strip():
                    # Limpiar formato (1.111,00 -> 1111.00)
                    try:
                        valor_limpio = str(ofertas_valor).replace('.', '').replace(',', '.')
                        precio_numerico = float(valor_limpio)
                        if precio_numerico > 0:
                            presenta_oferta = "Sí"
                    except:
                        precio_numerico = 0
                
                datos_tabla.append({
                    'nombre': nombre,
                    'presenta': presenta_oferta,
                    'precio': precio_numerico,
                    'ofertas_original': ofertas_valor
                })
            
            # 2. ORDENAR POR PRECIO (los que no presentan van al final)
            datos_tabla.sort(key=lambda x: (x['precio'] == 0, x['precio']))
            
            # 3. ASIGNAR ORDEN CLASIFICATORIO
            orden_clasif = 1
            for dato in datos_tabla:
                if dato['precio'] > 0:
                    dato['orden'] = str(orden_clasif)
                    orden_clasif += 1
                else:
                    dato['orden'] = "-"
            
            # 4. MOSTRAR TABLA POR LOG (PARA DEBUG)
            logger.debug(f"\n[DEBUG] 📋 TABLA FINAL:")
            logger.debug(f"{'Nombre':<40} {'¿Presenta?':<12} {'Importe':<15} {'Orden':<10}")
            logger.debug("-" * 80)
            for dato in datos_tabla:
                nombre_corto = dato['nombre'][:38] + ".." if len(dato['nombre']) > 40 else dato['nombre']
                importe_str = f"{dato['precio']:.2f}" if dato['precio'] > 0 else "0.00"
                logger.debug(f"{nombre_corto:<40} {dato['presenta']:<12} {importe_str:<15} {dato['orden']:<10}")
            logger.debug("-" * 80)
            
            # 5. CREAR TABLA EN WORD
            tabla = doc.add_table(rows=len(datos_tabla) + 1, cols=4)
            
            # Configurar bordes básicos
            if hasattr(self, '_configurar_tabla_basica'):
                self._configurar_tabla_basica(tabla)
            
            # 6. ENCABEZADOS
            encabezados = ["Nombre", "¿PRESENTA OFERTA?", "IMPORTE DE LA OFERTA", "ORDEN CLASIFICATORIO"]
            fila_enc = tabla.rows[0]
            for j, encabezado in enumerate(encabezados):
                celda = fila_enc.cells[j]
                celda.text = ""
                run = celda.paragraphs[0].add_run(encabezado)
                run.font.size = Pt(10)
                run.bold = True
            
            # 7. LLENAR DATOS
            for i, dato in enumerate(datos_tabla):
                fila = tabla.rows[i + 1]
                
                # Columna 0: Nombre
                celda = fila.cells[0]
                celda.text = ""
                run = celda.paragraphs[0].add_run(dato['nombre'])
                run.font.size = Pt(10)
                
                # Columna 1: ¿Presenta oferta?
                celda = fila.cells[1]
                celda.text = ""
                run = celda.paragraphs[0].add_run(dato['presenta'])
                run.font.size = Pt(10)
                
                # Columna 2: Importe
                importe_texto = f"{dato['precio']:.2f}" if dato['precio'] > 0 else "-"
                celda = fila.cells[2]
                celda.text = ""
                run = celda.paragraphs[0].add_run(importe_texto)
                run.font.size = Pt(10)
                
                # Columna 3: Orden
                celda = fila.cells[3]
                celda.text = ""
                run = celda.paragraphs[0].add_run(dato['orden'])
                run.font.size = Pt(10)
            
            # 8. INSERTAR TABLA EN EL DOCUMENTO
            if hasattr(self, '_insertar_tabla_despues_de_parrafo'):
                self._insertar_tabla_despues_de_parrafo(paragraph, tabla)
            
            # 9. LIMPIAR MARCADOR
            paragraph.text = paragraph.text.replace('@tabla-ofertas@', '')
            
            logger.debug(f"[DEBUG] ✅ Tabla insertada correctamente en Word")
            return datos_tabla
            
        except Exception as e:
            logger.error(f"[DEBUG] ❌ ERROR: {e}")
            import traceback
            logger.exception("Error completo:")
            paragraph.text = paragraph.text.replace('@tabla-ofertas@', '[Error procesando tabla]')
            return []

    def _insertar_tabla_anualidades(self, doc: Document, paragraph):
        """PASO 2: Agregar datos dinámicos de widgets"""
        try:
            logger.debug(f"[DEBUG] 📊 PASO 2: CREANDO TABLA CON DATOS DINÁMICOS")
            
            # 1. OBTENER DATOS DE WIDGETS
            ano_actual = self._obtener_valor_widget_directo('anoactual', '2025')
            ano_siguiente = self._obtener_valor_widget_directo('anosiguinte', '2026')
            
            anualidad1_sin_iva = self._obtener_valor_widget_directo('BaseAnualidad1', '0,00')
            anualidad1_iva = self._obtener_valor_widget_directo('IvaAnualidad1', '0,00')
            anualidad1_con_iva = self._obtener_valor_widget_directo('TotalAnualidad1', '0,00')
            
            anualidad2_sin_iva = self._obtener_valor_widget_directo('BaseAnualidad2', '0,00')
            anualidad2_iva = self._obtener_valor_widget_directo('IvaAnualidad2', '0,00')
            anualidad2_con_iva = self._obtener_valor_widget_directo('TotalAnualidad2', '0,00')
            
            logger.debug(f"[DEBUG] Datos obtenidos - Año actual: {ano_actual}, Anualidad1 sin IVA: {anualidad1_sin_iva}")
            
            # 2. DECIDIR SI INCLUIR SEGUNDA ANUALIDAD
            anualidad2_tiene_valor = self._tiene_valores_significativos(anualidad2_sin_iva, anualidad2_iva, anualidad2_con_iva)
            incluir_ano_siguiente = bool(ano_siguiente and str(ano_siguiente).strip()) and anualidad2_tiene_valor
            
            num_filas = 3 if incluir_ano_siguiente else 2  # Header + datos + total
            logger.debug(f"[DEBUG] Incluir año siguiente: {incluir_ano_siguiente}, Filas: {num_filas}")
            
            # 3. CREAR TABLA DINÁMICA
            tabla = doc.add_table(rows=num_filas + 1, cols=4)  # +1 para header
            
            # 4. HEADER
            tabla.cell(0, 0).text = "ANUALIDAD"
            tabla.cell(0, 1).text = "TOTAL SIN IVA"
            tabla.cell(0, 2).text = "IMPORTE IVA"
            tabla.cell(0, 3).text = "TOTAL CON IVA"
            
            # 5. FILA AÑO ACTUAL
            tabla.cell(1, 0).text = str(ano_actual)
            tabla.cell(1, 1).text = f"{anualidad1_sin_iva} €"
            tabla.cell(1, 2).text = f"{anualidad1_iva} €"
            tabla.cell(1, 3).text = f"{anualidad1_con_iva} €"
            
            # 6. FILA AÑO SIGUIENTE (SI EXISTE)
            fila_total_index = 2
            if incluir_ano_siguiente:
                tabla.cell(2, 0).text = str(ano_siguiente)
                tabla.cell(2, 1).text = f"{anualidad2_sin_iva} €"
                tabla.cell(2, 2).text = f"{anualidad2_iva} €"
                tabla.cell(2, 3).text = f"{anualidad2_con_iva} €"
                fila_total_index = 3
            
            # 7. FILA TOTAL
            total_sin_iva = self._sumar_importes(anualidad1_sin_iva, anualidad2_sin_iva if incluir_ano_siguiente else '0,00')
            total_iva = self._sumar_importes(anualidad1_iva, anualidad2_iva if incluir_ano_siguiente else '0,00')
            total_con_iva = self._sumar_importes(anualidad1_con_iva, anualidad2_con_iva if incluir_ano_siguiente else '0,00')
            
            tabla.cell(fila_total_index, 0).text = "TOTAL"
            tabla.cell(fila_total_index, 1).text = f"{total_sin_iva} €"
            tabla.cell(fila_total_index, 2).text = f"{total_iva} €"
            tabla.cell(fila_total_index, 3).text = f"{total_con_iva} €"
            
            # 8. CREAR TABLA CON TABULACIONES SOLO EN DATOS
            tabla_texto = f"""
ANUALIDAD            TOTAL SIN IVA           IMPORTE IVA             TOTAL CON IVA
{str(ano_actual)}\t\t{anualidad1_sin_iva} €\t\t{anualidad1_iva} €\t\t{anualidad1_con_iva} €"""
            
            if incluir_ano_siguiente:
                tabla_texto += f"""
{str(ano_siguiente)}\t\t{anualidad2_sin_iva} €\t\t{anualidad2_iva} €\t\t{anualidad2_con_iva} €"""
            
            tabla_texto += f"""
TOTAL\t\t{total_sin_iva} €\t\t{total_iva} €\t\t{total_con_iva} €
"""
            
            # 9. REEMPLAZAR MARCADOR Y APLICAR FORMATO ESPECÍFICO
            # Limpiar párrafo y agregar texto con formato Arial 7
            paragraph.clear()
            run = paragraph.add_run(tabla_texto.strip())
            
            # Aplicar formato: Arial 7, sin negrita
            from docx.shared import Pt
            run.font.name = 'Arial'
            run.font.size = Pt(7)
            run.font.bold = False
            
            logger.debug(f"[DEBUG] ✅ Tabla insertada y marcador limpiado exitosamente")
            
        except Exception as e:
            logger.error(f"[DEBUG] ❌ ERROR en tabla anualidades: {e}")
            import traceback
            logger.exception("Error completo:")
            paragraph.text = paragraph.text.replace('@tablaAnualidades@', '[Error procesando tabla anualidades]')

    def _obtener_valor_widget_directo(self, widget_name: str, default: str = '') -> str:
        """Obtener valor directamente del widget por nombre"""
        try:
            if hasattr(self, 'main_window') and hasattr(self.main_window, widget_name):
                widget = getattr(self.main_window, widget_name)
                
                # Para QDateEdit (años)
                if hasattr(widget, 'date'):
                    fecha = widget.date()
                    return str(fecha.year())
                
                # Para QDoubleSpinBox (importes)
                elif hasattr(widget, 'value'):
                    valor = widget.value()
                    # Convertir a formato español (1234.56 -> 1.234,56)
                    return f"{valor:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                
                # Para QLineEdit u otros
                elif hasattr(widget, 'text'):
                    return widget.text()
                
            return default
        except Exception as e:
            logger.debug(f"Error obteniendo widget {widget_name}: {e}")
            return default
            
    def _obtener_valor_widget(self, marcador: str, default: str = '') -> str:
        """Obtener valor de widget usando el sistema de sustitución existente"""
        try:
            # Usar el método de sustitución que ya existe
            texto_temp = marcador
            if hasattr(self, 'main_window'):
                # Llamar al método de sustitución existente
                resultado = self._sustituir_variables_en_texto(texto_temp, self.main_window)
                if resultado != marcador:  # Si se sustituyó
                    return resultado
            return default
        except:
            return default
            
    def _tiene_valores_significativos(self, *importes) -> bool:
        """Verificar si alguno de los importes es mayor que cero"""
        try:
            for importe in importes:
                if not importe or importe.strip() == '':
                    continue
                # Convertir formato español a float (1.234,56 -> 1234.56)
                limpio = str(importe).replace('.', '').replace(',', '.')
                valor = float(limpio)
                if valor > 0:
                    return True
            return False
        except:
            return False

    def _sumar_importes(self, importe1: str, importe2: str) -> str:
        """Sumar dos importes en formato español (1.234,56)"""
        try:
            # Convertir formato español a float
            def convertir_importe(importe_str):
                if not importe_str or importe_str.strip() == '':
                    return 0.0
                # Limpiar y convertir (1.234,56 -> 1234.56)
                limpio = str(importe_str).replace('.', '').replace(',', '.')
                return float(limpio)
            
            valor1 = convertir_importe(importe1)
            valor2 = convertir_importe(importe2)
            total = valor1 + valor2
            
            # Convertir de vuelta a formato español
            return f"{total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        except:
            return "0,00"

    def _configurar_tabla_basica(self, tabla):
        """Configurar estilo básico de tabla"""
        try:
            from docx.shared import Inches
            
            # Ancho de columnas
            for columna in tabla.columns:
                columna.width = Inches(1.5)
            
            # Bordes básicos
            for fila in tabla.rows:
                for celda in fila.cells:
                    tc = celda._element
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement("w:tcBorders")
                    
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:color'), '000000')
                        tcBorders.append(border)
                    
                    tcPr.append(tcBorders)
        except Exception:
            pass

    def _configurar_tabla_sin_bordes(self, tabla):
        """Configurar tabla sin bordes visibles"""
        try:
            from docx.shared import Inches
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            # Ancho de columnas
            for columna in tabla.columns:
                columna.width = Inches(1.5)
            
            # Sin bordes - configurar como transparentes/none
            for fila in tabla.rows:
                for celda in fila.cells:
                    tc = celda._element
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement("w:tcBorders")
                    
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f"w:{border_name}")
                        border.set(qn('w:val'), 'none')  # Sin bordes
                        tcBorders.append(border)
                    
                    tcPr.append(tcBorders)
        except Exception:
            pass

    def _insertar_tabla_despues_de_parrafo(self, paragraph, tabla):
        """Insertar tabla después de un párrafo específico"""
        try:
            paragraph_element = paragraph._element
            table_element = tabla._element
            paragraph_element.addnext(table_element)
        except Exception:
            pass
    def verificar_plantilla_disponible(self, nombre_plantilla: str) -> bool:
        """Verificar que una plantilla está disponible y accesible - ARREGLADO"""
        try:
            
            # Obtener ruta de la plantilla usando el sistema existente
            ruta_plantilla = self._obtener_ruta_plantilla(nombre_plantilla)
            
            if not ruta_plantilla:
                logger.error(f"[ControladorDocumentos] ❌ Plantilla no encontrada: {nombre_plantilla}")
                return False
            
            if not os.path.exists(ruta_plantilla):
                logger.error(f"[ControladorDocumentos] ❌ Archivo de plantilla no existe: {ruta_plantilla}")
                return False
            
            # Verificar que no está en uso
            try:
                with open(ruta_plantilla, 'rb') as test_file:
                    test_file.read(1)
                return True
            except PermissionError:
                logger.warning(f"[ControladorDocumentos] ❌ Plantilla en uso: {ruta_plantilla}")
                self._mostrar_error(f"La plantilla '{nombre_plantilla}' está siendo usada por Word.\nCierra Word y vuelve a intentar.")
                return False
                
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error verificando plantilla: {e}")
            return False
    def _procesar_paragraph_con_variables(self, paragraph, datos_json: Dict[str, Any]) -> set:
        """Procesar un párrafo sustituyendo variables @campo@ y \\@campo@ - VERSIÓN ROBUSTA"""
        variables_encontradas = set()
        
        try:
            if not paragraph or not paragraph.runs:
                return variables_encontradas
            
            texto_original = paragraph.text
            if not texto_original or not texto_original.strip():
                return variables_encontradas
            
            # Buscar tanto @campo@ como \\@campo@
            patron_variables = r'\\?@(\w+)@'
            variables_en_texto = re.findall(patron_variables, texto_original)
            variables_encontradas.update(variables_en_texto)
            
            if not variables_en_texto:
                return variables_encontradas
            
            texto_procesado = self._sustituir_variables_en_texto(texto_original, datos_json)
            
            if texto_original != texto_procesado:
                # Preservar formato del primer run
                formato_original = None
                if paragraph.runs and len(paragraph.runs) > 0:
                    try:
                        run_original = paragraph.runs[0]
                        formato_original = {
                            'font_name': run_original.font.name,
                            'font_size': run_original.font.size,
                            'bold': run_original.font.bold,
                            'italic': run_original.font.italic,
                            'underline': run_original.font.underline
                        }
                    except Exception:
                        formato_original = None
                
                # Limpiar párrafo de forma segura
                try:
                    paragraph.clear()
                    
                    # Crear nuevo run con texto procesado
                    new_run = paragraph.add_run(texto_procesado)
                    
                    # Restaurar formato si es posible
                    if formato_original:
                        try:
                            if formato_original.get('font_name'):
                                new_run.font.name = formato_original['font_name']
                            if formato_original.get('font_size'):
                                new_run.font.size = formato_original['font_size']
                            if formato_original.get('bold'):
                                new_run.font.bold = formato_original['bold']
                            if formato_original.get('italic'):
                                new_run.font.italic = formato_original['italic']
                            if formato_original.get('underline'):
                                new_run.font.underline = formato_original['underline']
                        except Exception:
                            pass  # Ignorar errores de formato
                            
                except Exception as e:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Error actualizando párrafo: {e}")
                    # Si falla la actualización, al menos registrar las variables encontradas
                    pass
        
        except Exception as e:
            logger.warning(f"[ControladorDocumentos] ⚠️ Error procesando párrafo: {e}")
        
        return variables_encontradas

    def _sustituir_variables_en_texto(self, texto: str, datos_json: Dict[str, Any]) -> str:
        patron = r'@(\w+)@'
        vars_encontradas = re.findall(patron, texto)
        if not vars_encontradas:
            return texto
        resultado = texto
        for var in vars_encontradas:
            if var in datos_json:
                val = datos_json[var]
                # Nuevo bloque para cadenas numéricas
                if isinstance(val, str):
                    # limpia espacios y prueba si es número
                    s = val.strip()
                    if re.fullmatch(r'\d+([.,]\d+)?', s):
                        try:
                            num = float(s.replace(',', '.'))
                            val_str = formatear_numero_espanol(num)
                        except:
                            val_str = s
                    else:
                        val_str = s
                elif isinstance(val, (int, float)):
                    val_str = formatear_numero_espanol(float(val))
                elif val is None:
                    val_str = ""
                else:
                    val_str = str(val)
                resultado = resultado.replace(f'@{var}@', val_str)
        return resultado

    # =================== MÉTODOS PARA CARTAS INDIVIDUALES ===================

    def _preparar_datos_carta_empresa(self, contract_data, empresa, indice_empresa):
        """Preparar datos específicos para carta de invitación"""
        try:
            datos_carta = self._preparar_datos_para_sustitucion(contract_data)
            
            # Datos específicos de la empresa
            datos_empresa = [
                empresa.get('nombre', ''),
                empresa.get('nif', ''),
                empresa.get('contacto', ''),
                empresa.get('email', '')
            ]
            
            for i, valor in enumerate(datos_empresa):
                datos_carta[f'tabla{i}'] = valor
            
            # Variables adicionales para cartas de invitación
            datos_carta.update({
                'numeroEmpresa': str(indice_empresa + 1),
                'empresaActual': empresa.get('nombre', ''),
                'horaDeApertura': contract_data.get('horaDeApertura', '14:00'),
                'diaDeApertura': contract_data.get('diaDeApertura', 'por determinar'),
                'mailDeRecepcion': contract_data.get('mailDeRecepcion', 'licitaciones@adif.es'),
                'consultasAdministrativas': contract_data.get('consultasAdministrativas', 'consultas.admin@adif.es'),
                'consultasTecnicas': contract_data.get('consultasTecnicas', 'consultas.tecnicas@adif.es')
            })
            
            return datos_carta
        except Exception:
            return contract_data

    def generar_cartas_adjudicacion(self):
        """Generar cartas de adjudicación: solo una adjudicataria, resto no adjudicatarios"""
        try:
            if not self._validar_contrato_seleccionado():
                logger.debug("[DEBUG] Validación de contrato fallida")
                return
            
            contract_data = self._obtener_datos_contrato_actual()
            #logger.info(f"[DEBUG] contract_data keys: {list(contract_data.keys()) if contract_data else 'None'}")
            if not contract_data:
                logger.debug("[DEBUG] No se pudieron obtener los datos del contrato")
                return self._mostrar_error("No se pudieron obtener los datos del contrato")
            
            empresas_lista = self._obtener_empresas_lista(contract_data)
            logger.debug(f"[DEBUG] Total empresas para cartas adjudicación: {len(empresas_lista)}")
            if not empresas_lista:
                logger.debug("[DEBUG] No hay empresas para generar cartas de adjudicación")
                return self._mostrar_error("No hay empresas para generar cartas de adjudicación")
            
            plantilla_adjudicacion = self._obtener_ruta_plantilla('plantilla_cartas_adjudicacion.docx')
            plantilla_noadjudicacion = self._obtener_ruta_plantilla('plantilla_cartas_noadjudicacion.docx')
            
            logger.debug(f"[DEBUG] plantilla_adjudicacion: {plantilla_adjudicacion}")
            logger.debug(f"[DEBUG] plantilla_noadjudicacion: {plantilla_noadjudicacion}")
            if not plantilla_adjudicacion or not plantilla_noadjudicacion:
                logger.debug("[DEBUG] No se encontraron las plantillas de adjudicación/no adjudicación")
                return self._mostrar_error("No se encontraron las plantillas de adjudicación/no adjudicación")
            
            carpeta_contrato = self._obtener_carpeta_con_gestor_unificado(contract_data)
            logger.debug(f"[DEBUG] carpeta_contrato: {carpeta_contrato}")
            if not carpeta_contrato:
                logger.debug("[DEBUG] No se pudo crear/obtener la carpeta del contrato")
                return self._mostrar_error("No se pudo crear/obtener la carpeta del contrato")
            
            subcarpeta_cartas = self._determinar_subcarpeta_por_tipo_documento('generar_cartas_adjudicacion')
            directorio_cartas = os.path.join(carpeta_contrato, subcarpeta_cartas)
            os.makedirs(directorio_cartas, exist_ok=True)
            
            empresa_adjudicataria = contract_data.get('empresaAdjudicada', '').strip().lower()
            logger.debug(f"[DEBUG] 🏆 empresaAdjudicataria (comparación): '{empresa_adjudicataria}'")
            
            cartas_generadas = []
            
            for i, empresa in enumerate(empresas_lista):
                nombre_empresa = empresa.get('nombre', f'Empresa_{i+1}')
                nombre_empresa_limpio = self._limpiar_nombre_para_archivo(nombre_empresa)
                logger.debug(f"[DEBUG] Empresa {i+1}: '{nombre_empresa}' (comparando con adjudicataria '{empresa_adjudicataria}')")
                logger.debug(f"[DEBUG] Empresa dict: {empresa}")
                # Solo una adjudicataria, resto no adjudicatarios
                if nombre_empresa.strip().lower() == empresa_adjudicataria and empresa_adjudicataria:
                    logger.debug(f"[DEBUG] --> adjudicataria detectada: '{nombre_empresa}'")
                    plantilla_usar = plantilla_adjudicacion
                    nombre_archivo = f"Carta_Adjudicatario_{i+1:02d}_{nombre_empresa_limpio}.docx"
                    tipo_carta = "ADJUDICATARIA"
                    es_adjudicataria = True
                else:
                    logger.debug(f"[DEBUG] --> NO adjudicataria: '{nombre_empresa}'")
                    plantilla_usar = plantilla_noadjudicacion
                    nombre_archivo = f"Carta_No_Adjudicatario_{i+1:02d}_{nombre_empresa_limpio}.docx"
                    tipo_carta = "NO ADJUDICATARIA"
                    es_adjudicataria = False
                
                archivo_salida = os.path.join(directorio_cartas, nombre_archivo)
                
                logger.debug(f"[DEBUG] Generando carta {tipo_carta} para: {nombre_empresa} -> {archivo_salida}")
                
                datos_carta = self._preparar_datos_carta_adjudicacion(contract_data, empresa, i, es_adjudicataria)
                
                if self._generar_carta_individual(plantilla_usar, archivo_salida, datos_carta):
                    logger.debug(f"[DEBUG] Carta generada correctamente: {archivo_salida}")
                    cartas_generadas.append({
                        'archivo': archivo_salida,
                        'empresa': nombre_empresa,
                        'numero': i + 1,
                        'tipo': tipo_carta
                    })
                else:
                    logger.error(f"[DEBUG] Error generando carta para: {nombre_empresa}")
            
            logger.debug(f"[DEBUG] Total cartas generadas: {len(cartas_generadas)}")
            self._mostrar_resultado_cartas(cartas_generadas, contract_data, "Cartas de Adjudicación")
            if cartas_generadas:
                self._abrir_cartas_generadas(cartas_generadas)
                
                # 🆕 NOTIFICAR AL CRONOGRAMA QUE SE GENERARON LAS CARTAS DE ADJUDICACIÓN
                self._actualizar_fase_en_generacion('cartas_adjudicacion')
            
        except Exception as e:
            logger.error(f"[DEBUG] Error generando cartas de adjudicación: {str(e)}")
            self._mostrar_error(f"Error generando cartas de adjudicación: {str(e)}")
    def _preparar_datos_carta_adjudicacion(self, contract_data, empresa, indice_empresa, es_adjudicataria=False):
        """Preparar datos individualizados para carta de adjudicación/no adjudicataria"""
        try:
            datos_carta = self._preparar_datos_para_sustitucion(contract_data)
            # Añadir datos específicos de la empresa
            datos_empresa = [
                empresa.get('nombre', ''),
                empresa.get('nif', ''),
                empresa.get('contacto', ''),
                empresa.get('email', '')
            ]
            for i, valor in enumerate(datos_empresa):
                datos_carta[f'tabla{i}'] = valor
            # Añadir tipo de carta adjudicataria/no adjudicataria
            datos_carta['esAdjudicataria'] = 'SI' if es_adjudicataria else 'NO'
            datos_carta['tipoResultado'] = 'ADJUDICACIÓN' if es_adjudicataria else 'NO ADJUDICACIÓN'
            datos_carta['mensajeResultado'] = (
                'Ha resultado ADJUDICATARIA del contrato' if es_adjudicataria
                else 'NO ha resultado adjudicataria del contrato'
            )
            # Añadir datos comunes (sin fechaActual, ya viene en contract_data)
            datos_carta.update({
                'numeroEmpresa': str(indice_empresa + 1),
                'empresaActual': empresa.get('nombre', ''),
                'cifEmpresa': empresa.get('nif', ''),
                'contactoEmpresa': empresa.get('contacto', ''),
                'emailEmpresa': empresa.get('email', ''),
            })
            # Corregir conversión de oferta a float (acepta formato con puntos y comas)
            ofertas_valor = empresa.get('ofertas', '')
            precio_empresa_actual = "0.00"
            if ofertas_valor:
                try:
                    ofertas_valor_limpio = ofertas_valor.replace('.', '').replace(',', '.')
                    precio_empresa_actual = f"{float(ofertas_valor_limpio):.2f}"
                except Exception as e:
                    logger.error(f"[DEBUG] Error convirtiendo oferta: {ofertas_valor} -> {e}")
                    precio_empresa_actual = "0.00"
            datos_carta['precioEmpresaActual'] = precio_empresa_actual
            return datos_carta
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error preparando datos carta adjudicación: {e}")
            return contract_data

    def _generar_carta_individual(self, ruta_plantilla, archivo_salida, datos_carta):
        """Generar una carta individual sustituyendo variables en plantilla"""
        try:
            doc = Document(ruta_plantilla)
            # Procesar párrafos principales
            for paragraph in doc.paragraphs:
                self._procesar_paragraph_con_variables(paragraph, datos_carta)
            # Procesar tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._procesar_paragraph_con_variables(paragraph, datos_carta)
            # Procesar headers y footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self._procesar_paragraph_con_variables(paragraph, datos_carta)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self._procesar_paragraph_con_variables(paragraph, datos_carta)
            doc.save(archivo_salida)
            return True
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error generando carta individual: {e}")
            return False

    # =================== FUNCIONES DE VALIDACIÓN Y COMPROBACIÓN ===================

    def _validar_campos_y_fechas(self, contract_data, nombre_plantilla):
        """Validar campos vacíos y fechas fuera de rango"""
        try:
            # Detectar variables en la plantilla
            ruta_plantilla = self._obtener_ruta_plantilla(nombre_plantilla)
            if not ruta_plantilla:
                return True, []  # Si no encuentra plantilla, permitir continuar
                
            variables_plantilla = self._detectar_variables_en_plantilla(ruta_plantilla)
            
            # Preparar datos completos
            datos_completos = self._preparar_datos_para_sustitucion(contract_data)
            
            # Validar campos vacíos
            campos_vacios = []
            fechas_incorrectas = []
            
            for variable in variables_plantilla:
                # Omitir tablas especiales y campos que empiecen con tabla-
                if (variable.startswith('tabla-') or variable == 'tablaAnualidades' or
                    variable.lower() in ['tabla0', 'tabla1', 'tabla2', 'tabla3']):
                    continue
                    
                if variable not in datos_completos:
                    campos_vacios.append(f"{variable} (no existe)")
                else:
                    valor = datos_completos[variable]
                    
                    # Verificar campos vacíos
                    if not valor or str(valor).strip() in ["", "0", "0.0", "0.00", "---------------------"]:
                        campos_vacios.append(f"{variable} (vacío)")
                    
                    # Verificar fechas
                    if self._es_campo_fecha(variable):
                        fecha_valida = self._validar_fecha_rango(valor)
                        if not fecha_valida:
                            fechas_incorrectas.append(f"{variable} (fecha no tiene sentido: {valor})")
            
            return len(campos_vacios) == 0 and len(fechas_incorrectas) == 0, campos_vacios + fechas_incorrectas
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error validando: {e}")
            return True, []  # En caso de error, permitir continuar

    def _es_campo_fecha(self, nombre_campo):
        """Determinar si un campo es una fecha"""
        campos_fecha = [
            'fecha', 'fechaInicio', 'fechaFin', 'fechaAdjudicacion', 'fechaRecepcion',
            'fechaReplanteo', 'fechaFinal', 'fechaCreacion', 'fechaGeneracion',
            'fechaLegalFin', 'fechaProyecto', 'fechaInforme', 'diaDeApertura'
        ]
        es_fecha = any(campo_fecha.lower() in nombre_campo.lower() for campo_fecha in campos_fecha)
        #logger.info(f"[DEBUG] Campo '{nombre_campo}' es fecha: {es_fecha}")
        return es_fecha

    def _validar_fecha_rango(self, valor_fecha):
        """Validar que la fecha esté en rango de ±6 meses"""
        try:
            from datetime import datetime, timedelta
            import re
            
            if not valor_fecha or not str(valor_fecha).strip():
                return True  # Fecha vacía se considera válida para esta validación
            
            fecha_str = str(valor_fecha).strip()
            fecha_obj = None
            
            # Intentar parsear diferentes formatos de fecha
            formatos = [
                "%Y-%m-%d",      # 2024-01-15
                "%d/%m/%Y",      # 15/01/2024
                "%d-%m-%Y",      # 15-01-2024
                "%Y/%m/%d",      # 2024/01/15
            ]
            
            for formato in formatos:
                try:
                    fecha_obj = datetime.strptime(fecha_str, formato)
                    logger.debug(f"[DEBUG] Fecha parseada: {fecha_str} -> {fecha_obj} usando formato {formato}")
                    break
                except ValueError:
                    continue
            
            if not fecha_obj:
                logger.warning(f"[DEBUG] No se pudo parsear fecha: {fecha_str}")
                return True  # Si no puede parsear, considera válida
            
            # Verificar rango de ±6 meses
            hoy = datetime.now()
            limite_anterior = hoy - timedelta(days=180)  # 6 meses atrás
            limite_posterior = hoy + timedelta(days=180)  # 6 meses adelante
            
            es_valida = limite_anterior <= fecha_obj <= limite_posterior
            logger.debug(f"[DEBUG] Validación fecha: {fecha_obj} | Rango: {limite_anterior} - {limite_posterior} | Válida: {es_valida}")
            
            return es_valida
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error validando fecha {valor_fecha}: {e}")
            return True  # En caso de error, considerar válida

    def _mostrar_popup_validacion(self, problemas_encontrados, nombre_documento):
        """Mostrar popup con problemas de validación"""
        try:
            from PyQt5.QtWidgets import QMessageBox
            
            # Separar campos vacíos y fechas incorrectas
            campos_vacios = [p for p in problemas_encontrados if 'vacío' in p or 'no existe' in p]
            fechas_incorrectas = [p for p in problemas_encontrados if p not in campos_vacios]
            
            # Preparar mensaje
            mensaje = f"<b>Se encontraron problemas en '{nombre_documento}':</b><br><br>"
            
            if campos_vacios:
                mensaje += f"<b>📝 Campos vacíos ({len(campos_vacios)}):</b><br>"
                for campo in campos_vacios[:10]:  # Máximo 10
                    mensaje += f"• {campo}<br>"
                if len(campos_vacios) > 10:
                    mensaje += f"• ... y {len(campos_vacios) - 10} más<br>"
                mensaje += "<br>"
            
            if fechas_incorrectas:
                mensaje += f"<b>📅 Fechas fuera de rango (±6 meses) ({len(fechas_incorrectas)}):</b><br>"
                for fecha in fechas_incorrectas[:5]:  # Máximo 5
                    mensaje += f"• {fecha}<br>"
                if len(fechas_incorrectas) > 5:
                    mensaje += f"• ... y {len(fechas_incorrectas) - 5} más<br>"
                mensaje += "<br>"
            
            mensaje += "<b>¿Deseas generar el documento de todas formas?</b>"
            
            msg = QMessageBox(self.main_window)
            msg.setIcon(QMessageBox.Warning)
            msg.setWindowTitle("⚠️ Problemas de Validación")
            msg.setText(mensaje)
            msg.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg.setDefaultButton(QMessageBox.No)
            
            # Personalizar botones
            yes_button = msg.button(QMessageBox.Yes)
            no_button = msg.button(QMessageBox.No)
            yes_button.setText("✅ Generar de todas formas")
            no_button.setText("❌ Cancelar y revisar")
            
            resultado = msg.exec_()
            return resultado == QMessageBox.Yes
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error mostrando popup validación: {e}")
            return True  # En caso de error, permitir continuar

    # =================== FUNCIONES COMPROBAR_GENERAR ===================

    def comprobar_generar_acta_inicio(self):
        """Comprobar y generar acta de inicio"""
        if self._comprobar_y_ejecutar('plantilla_acta_inicio.docx', self.generar_acta_inicio):
            return True
        return False

    def comprobar_generar_cartas_invitacion(self):
        """Comprobar y generar cartas de invitación"""
        if self._comprobar_y_ejecutar('plantilla_cartas_invitacion.docx', self.generar_cartas_invitacion):
            return True
        return False

    def comprobar_generar_acta_adjudicacion(self):
        """Comprobar y generar acta de adjudicación"""
        if self._comprobar_y_ejecutar('plantilla_acta_adjudicacion.docx', self.generar_acta_adjudicacion):
            return True
        return False

    def comprobar_generar_cartas_adjudicacion(self):
        """Comprobar y generar cartas de adjudicación"""
        if self._comprobar_y_ejecutar('plantilla_cartas_adjudicacion.docx', self.generar_cartas_adjudicacion):
            return True
        return False

    def comprobar_generar_acta_liquidacion(self):
        """Comprobar y generar acta de liquidación"""
        if self._comprobar_y_ejecutar('plantilla_acta_finalizacion.docx', self.generar_acta_liquidacion):
            return True
        return False

    def comprobar_generar_acta_replanteo(self):
        """Comprobar y generar acta de replanteo"""
        if self._comprobar_y_ejecutar('plantilla_acta_replanteo.docx', self.generar_acta_replanteo):
            return True
        return False

    def comprobar_generar_acta_recepcion(self):
        """Comprobar y generar acta de recepción"""
        if self._comprobar_y_ejecutar('plantilla_acta_recepcion.docx', self.generar_acta_recepcion):
            return True
        return False

    def comprobar_generar_nombramiento_director(self):
        """Comprobar y generar nombramiento director"""
        if self._comprobar_y_ejecutar('Modelo_director_obra_ajuste.docx', self.generar_nombramiento_director):
            return True
        return False

    def comprobar_generar_contrato(self):
        """Comprobar y generar contrato CON selección dinámica"""
        # Obtener plantilla dinámica para validación
        nombre_plantilla = self._obtener_nombre_plantilla_dinamico('generar_contrato')
        if self._comprobar_y_ejecutar(nombre_plantilla, self.generar_contrato):
            return True
        return False

    def _comprobar_y_ejecutar(self, nombre_plantilla, funcion_generar):
        """Función auxiliar para comprobar y ejecutar generación"""
        try:
            # Validar contrato seleccionado
            if not self._validar_contrato_seleccionado():
                return False
            
            # Obtener datos del contrato
            contract_data = self._obtener_datos_contrato_actual()
            if not contract_data:
                self._mostrar_error("No se pudieron obtener los datos del contrato")
                return False
            
            # Validar campos y fechas
            validacion_ok, problemas = self._validar_campos_y_fechas(contract_data, nombre_plantilla)
            
            # Si hay problemas, mostrar popup
            if not validacion_ok:
                nombre_documento = nombre_plantilla.replace('plantilla_', '').replace('.docx', '').replace('_', ' ').title()
                continuar = self._mostrar_popup_validacion(problemas, nombre_documento)
                
                if not continuar:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Generación cancelada por validación")
                    return False
            
            # Ejecutar función de generación
            logger.info(f"[ControladorDocumentos] ✅ Validación pasada, generando documento")
            funcion_generar()
            return True
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error en comprobar y ejecutar: {e}")
            self._mostrar_error(f"Error en validación: {str(e)}")
            return False

    def _validar_contrato_seleccionado(self):
        """Validar que hay un contrato seleccionado"""
        try:
            if not hasattr(self.main_window, 'comboBox') or not self.main_window.comboBox:
                self._mostrar_error("No se encontró selector de contratos")
                return False
            
            texto_seleccionado = self.main_window.comboBox.currentText()
            if not texto_seleccionado or texto_seleccionado.strip() == "":
                self._mostrar_error("No hay ningún contrato seleccionado")
                return False
                
            return True
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error validando contrato: {e}")
            return False

    def _obtener_datos_contrato_actual(self):
        """Obtener datos del contrato actualmente seleccionado"""
        try:
            if not hasattr(self.main_window, 'controlador_json'):
                return None
                
            # Obtener nombre del contrato seleccionado
            if not hasattr(self.main_window, 'comboBox') or not self.main_window.comboBox:
                return None
                
            nombre_contrato = self.main_window.comboBox.currentText()
            if not nombre_contrato or nombre_contrato.strip() == "":
                return None
            
            # 🆕 NUEVO: Guardar nombre del contrato para tracking
            self.contract_name = nombre_contrato
            
            # Leer datos completos del contrato
            datos_contrato = self.main_window.controlador_json.leer_contrato_completo(nombre_contrato)
            return datos_contrato
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error obteniendo datos contrato: {e}")
            return None

    def _mostrar_error(self, mensaje):
        """Mostrar mensaje de error"""
        try:
            from PyQt5.QtWidgets import QMessageBox
            QMessageBox.critical(self.main_window, "❌ Error", mensaje)
        except Exception as e:
            logger.error(f"[ControladorDocumentos] ❌ Error mostrando mensaje: {e}")

    # =================== FUNCIONES DE GENERACIÓN DE RESUMEN DOCX ===================
    # Funciones movidas desde controlador_resumen.py para unificar el código de documentos

    def generar_fichero_resumen(self, nombre_contrato: str, datos_contrato: dict) -> str:
        """Generar fichero completo de resumen del contrato en formato Word"""
        try:
            logger.info(f"[ControladorDocumentos] Iniciando generación de fichero para: {nombre_contrato}")
            
            # Crear tracker
            from .controlador_resumen import TrackerDocumentos
            tracker = TrackerDocumentos()
            
            logger.info("[ControladorDocumentos] Obteniendo resumen de documentos...")
            resumen_docs = tracker.obtener_resumen_contrato(nombre_contrato)
            
            # Obtener datos de firmas para la tabla de seguimiento
            logger.info("[ControladorDocumentos] Obteniendo datos de firmas...")
            datos_firmas = {}
            firmantes_unicos = []
            if hasattr(self.main_window, 'controlador_resumen'):
                try:
                    # Escanear firmas para obtener los datos actualizados
                    self.main_window.controlador_resumen._escanear_y_actualizar_tabla_firmas(nombre_contrato)
                    # Los datos quedan guardados internamente, necesitamos extraerlos
                    datos_firmas = getattr(self.main_window.controlador_resumen, '_datos_firmas_cache', {})
                    firmantes_unicos = getattr(self.main_window.controlador_resumen, '_firmantes_unicos_cache', [])
                    logger.info(f"[ControladorDocumentos] Datos de firmas obtenidos: {len(datos_firmas)} fases, {len(firmantes_unicos)} firmantes")
                except Exception as e:
                    logger.warning(f"[ControladorDocumentos] ⚠️ Error obteniendo datos de firmas: {e}")
            
            # Crear documento Word
            logger.info("[ControladorDocumentos] Creando documento Word...")
            documento_word = self._crear_contenido_fichero_resumen(
                nombre_contrato, datos_contrato, resumen_docs, tracker, datos_firmas, firmantes_unicos
            )
            
            # Guardar archivo Word
            logger.info("[ControladorDocumentos] Guardando archivo...")
            ruta_archivo = self._guardar_fichero_resumen(nombre_contrato, documento_word)
            
            logger.info(f"[ControladorDocumentos] ✅ Fichero generado exitosamente: {ruta_archivo}")
            return ruta_archivo
            
        except ImportError as e:
            error_msg = f"Error de importación: {e}"
            logger.error(f"[ControladorDocumentos] ❌ {error_msg}")
            raise ImportError(error_msg)
        except Exception as e:
            error_msg = f"Error generando fichero: {e}"
            logger.error(f"[ControladorDocumentos] ❌ {error_msg}")
            raise Exception(error_msg)
    
    def _crear_contenido_fichero_resumen(self, nombre_contrato: str, datos_contrato: dict, 
                                       resumen_docs: dict, tracker, datos_firmas: dict = None, firmantes_unicos: list = None) -> str:
        """Crear el documento Word del fichero resumen con cronograma y datos completos"""
        try:
            from datetime import datetime
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.shared import OxmlElement, qn
        except ImportError as e:
            raise ImportError(f"Error importando módulos de docx: {e}. Asegúrate de que python-docx esté instalado.")
        
        # Crear documento Word
        doc = Document()
        
        # Configurar márgenes del documento
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ENCABEZADO PRINCIPAL
        titulo = doc.add_heading('RESUMEN DEL CONTRATO', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del contrato y fecha
        info_header = doc.add_paragraph()
        info_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info_header.add_run(f'{nombre_contrato}')
        run.font.size = Pt(14)
        run.font.bold = True
        
        fecha_para = doc.add_paragraph()
        fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_run = fecha_para.add_run(f'Generado el {datetime.now().strftime("%d/%m/%Y a las %H:%M")}')
        fecha_run.font.size = Pt(11)
        fecha_run.font.italic = True
        
        # Línea separadora
        doc.add_paragraph('_' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # SECCIÓN 1: DATOS PRINCIPALES DEL CONTRATO (AMPLIADO)
        doc.add_heading('📋 DATOS PRINCIPALES DEL CONTRATO', level=1)
        
        # Datos básicos del contrato
        campos_principales = [
            ('nombreObra', 'Nombre de la Obra'),
            ('tipoActuacion', 'Tipo de Actuación'),
            ('numeroExpediente', 'Número de Expediente'),
            ('objeto', 'Objeto del Contrato'),
            ('justificacion', 'Justificación'),
            ('insuficiencia', 'Insuficiencia de Medios'),
            ('plazoEjecucion', 'Plazo de Ejecución (meses)'),
            ('fechaCreacion', 'Fecha de Creación'),
            ('basePresupuesto', 'Base Presupuesto'),
            ('ivaPresupuestoBase', 'IVA'),
            ('totalPresupuestoBase', 'Total Presupuesto'),
            ('organoContratacion2', 'Órgano de Contratación'),
            ('OrganoSolicitaOfertas', 'Órgano que Solicita Ofertas'),
            ('justificacionLimites', 'Justificación de Límites'),
            ('regimenPagos', 'Régimen de Pagos')
        ]
        
        # Crear tabla de 2 columnas para datos principales
        tabla_datos = doc.add_table(rows=len(campos_principales), cols=2)
        tabla_datos.style = 'Table Grid'
        
        for i, (campo, etiqueta) in enumerate(campos_principales):
            valor = datos_contrato.get(campo, 'No especificado')
            if campo in ['basePresupuesto', 'ivaPresupuestoBase', 'totalPresupuestoBase']:
                try:
                    if isinstance(valor, str):
                        valor = valor.replace(',', '.')
                    valor_num = float(valor) if valor and valor != 'No especificado' else 0
                    valor = f"{valor_num:,.2f} €"
                except:
                    valor = str(valor) if valor else 'No especificado'
            
            # Celda de etiqueta
            celda_etiqueta = tabla_datos.cell(i, 0)
            celda_etiqueta.text = etiqueta
            run = celda_etiqueta.paragraphs[0].runs[0]
            run.font.bold = True
            
            # Celda de valor
            celda_valor = tabla_datos.cell(i, 1)
            celda_valor.text = str(valor) if valor else 'No especificado'
        
        doc.add_paragraph()  # Espaciado
        
        # SECCIÓN 2: CRONOGRAMA DE FASES DEL PROYECTO
        doc.add_heading('📅 CRONOGRAMA DE FASES DEL PROYECTO', level=1)
        
        # Insertar cronograma de fases
        self._insertar_cronograma_en_word(doc, nombre_contrato)
        
        doc.add_paragraph()  # Espaciado
        
        # SECCIÓN 3: EMPRESAS PARTICIPANTES
        doc.add_heading('🏢 EMPRESAS PARTICIPANTES', level=1)
        
        empresas = datos_contrato.get('empresas', [])
        if empresas and isinstance(empresas, list):
            # Crear tabla para empresas
            tabla_empresas = doc.add_table(rows=1, cols=4)
            tabla_empresas.style = 'Table Grid'
            
            # Encabezados
            encabezados_empresas = ['Empresa', 'NIF/CIF', 'Email', 'Contacto']
            for i, encabezado in enumerate(encabezados_empresas):
                celda = tabla_empresas.cell(0, i)
                celda.text = encabezado
                run = celda.paragraphs[0].runs[0]
                run.font.bold = True
            
            # Agregar empresas
            for empresa in empresas:
                if isinstance(empresa, dict):
                    row_cells = tabla_empresas.add_row().cells
                    row_cells[0].text = empresa.get('nombre', 'Sin nombre')
                    row_cells[1].text = empresa.get('nif', 'Sin NIF')
                    row_cells[2].text = empresa.get('email', 'Sin email')
                    row_cells[3].text = empresa.get('contacto', 'Sin contacto')
        else:
            para_sin_empresas = doc.add_paragraph("🏢 No hay empresas registradas para este contrato")
            para_sin_empresas.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para_sin_empresas.runs[0]
            run.font.italic = True
        
        doc.add_paragraph()  # Espaciado
        
        # SECCIÓN 4: INFORMACIÓN DE FECHAS IMPORTANTES
        doc.add_heading('📅 FECHAS IMPORTANTES', level=1)
        
        campos_fechas = [
            ('fechaCreacion', 'Fecha de Creación'),
            ('fechaAdjudicacion', 'Fecha de Adjudicación'),
            ('fechaContrato', 'Fecha del Contrato'),
            ('fechaProyecto', 'Fecha de Inicio del Proyecto'),
            ('fechaReplanteo', 'Fecha de Replanteo'),
            ('fechaRecepcion', 'Fecha de Recepción'),
            ('fechaFinal', 'Fecha Final'),
            ('diaApertura', 'Día de Apertura'),
            ('horaDeApertura', 'Hora de Apertura')
        ]
        
        # Crear tabla para fechas
        tabla_fechas = doc.add_table(rows=len(campos_fechas), cols=2)
        tabla_fechas.style = 'Table Grid'
        
        for i, (campo, etiqueta) in enumerate(campos_fechas):
            valor = datos_contrato.get(campo, 'No especificado')
            
            # Celda de etiqueta
            celda_etiqueta = tabla_fechas.cell(i, 0)
            celda_etiqueta.text = etiqueta
            run = celda_etiqueta.paragraphs[0].runs[0]
            run.font.bold = True
            
            # Celda de valor
            celda_valor = tabla_fechas.cell(i, 1)
            celda_valor.text = str(valor) if valor else 'No especificado'
        
        doc.add_paragraph()  # Espaciado
        
        # SECCIÓN 5: HISTORIAL DE DOCUMENTOS (SIMPLIFICADO)
        doc.add_heading('📄 HISTORIAL DE DOCUMENTOS', level=1)
        
        documentos = tracker.obtener_documentos_contrato(nombre_contrato)
        if documentos:
            from .controlador_resumen import TipoDocumento
            iconos_docs = {
                TipoDocumento.INVITACION: "📧",
                TipoDocumento.ADJUDICACION: "🏆",
                TipoDocumento.ACTA_INICIO: "🚀",
                TipoDocumento.ACTA_REPLANTEO: "📐",
                TipoDocumento.ACTA_RECEPCION: "✅",
                TipoDocumento.ACTA_FINALIZACION: "🏁",
                TipoDocumento.LIQUIDACION: "💰",
                TipoDocumento.CONTRATO: "📋",
                TipoDocumento.OTRO: "📄"
            }
            
            # Lista simple de documentos sin columnas complejas
            for doc_item in sorted(documentos, key=lambda d: d.fecha_generacion, reverse=True):
                para_doc = doc.add_paragraph()
                icono = iconos_docs.get(doc_item.tipo, "📄")
                fecha_formato = doc_item.fecha_generacion.strftime('%d/%m/%Y %H:%M')
                
                # Documento con fecha y estado
                run_doc = para_doc.add_run(f"{icono} {doc_item.nombre}")
                run_doc.font.bold = True
                para_doc.add_run(f" - {fecha_formato} - {doc_item.estado.value.upper()}")
        else:
            para_sin_docs = doc.add_paragraph("📄 No hay documentos generados para este contrato")
            para_sin_docs.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para_sin_docs.runs[0]
            run.font.italic = True
        
        # FOOTER
        doc.add_page_break()  # Nueva página para footer si es necesario
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"Generado por ADIF - Generador de Actas | {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        
        footer_para2 = doc.add_paragraph()
        footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run2 = footer_para2.add_run("Este resumen incluye todos los datos disponibles del contrato en el momento de la generación")
        footer_run2.font.size = Pt(9)
        footer_run2.font.italic = True
        
        # SECCIÓN: TABLA DE SEGUIMIENTO DE FIRMAS
        if datos_firmas and firmantes_unicos:
            self._insertar_tabla_seguimiento_en_word(doc, datos_firmas, firmantes_unicos)
        
        return doc
    
    def _insertar_cronograma_en_word(self, doc, nombre_contrato: str):
        """Insertar cronograma de fases del proyecto en el documento Word"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Obtener datos de fases
            if hasattr(self.main_window, 'controlador_fases'):
                datos_fases = self.main_window.controlador_fases.obtener_datos_fases_para_resumen(nombre_contrato)
                
                if datos_fases:
                    # Crear tabla para el cronograma
                    tabla_cronograma = doc.add_table(rows=1, cols=3)
                    tabla_cronograma.style = 'Table Grid'
                    
                    # Encabezados
                    encabezados = ['Fase del Proyecto', 'Fecha Generación', 'Fecha Firma']
                    for i, encabezado in enumerate(encabezados):
                        celda = tabla_cronograma.cell(0, i)
                        celda.text = encabezado
                        run = celda.paragraphs[0].runs[0]
                        run.font.bold = True
                        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Mapeo de nombres de fases
                    nombres_fases = {
                        'creacion': '📋 Creación Proyecto',
                        'inicio': '🚀 Inicio',
                        'cartas': '📧 Cartas Invitación',
                        'adjudicacion': '🏆 Adjudicación',
                        'firma_contrato': '✍️ Firma Contrato',
                        'replanteo': '📐 Replanteo',
                        'actuacion': '⚡ Actuación',
                        'recepcion': '✅ Recepción',
                        'finalizacion': '🏁 Finalización'
                    }
                    
                    # Agregar fila para cada fase
                    for fase_key, fase_data in datos_fases.items():
                        row_cells = tabla_cronograma.add_row().cells
                        
                        # Nombre de la fase
                        nombre_fase = nombres_fases.get(fase_key, fase_key.title())
                        row_cells[0].text = nombre_fase
                        
                        # Fecha de generación
                        generado = fase_data.get('generado')
                        if generado:
                            row_cells[1].text = str(generado)
                            # Poner en verde si está generado
                            run_gen = row_cells[1].paragraphs[0].runs[0]
                            run_gen.font.bold = True
                        else:
                            row_cells[1].text = "Pendiente"
                            run_gen = row_cells[1].paragraphs[0].runs[0]
                            run_gen.font.italic = True
                        
                        # Fecha de firma
                        firmado = fase_data.get('firmado')
                        if firmado:
                            row_cells[2].text = str(firmado)
                            # Poner en verde si está firmado
                            run_firm = row_cells[2].paragraphs[0].runs[0]
                            run_firm.font.bold = True
                        else:
                            row_cells[2].text = "Pendiente"
                            run_firm = row_cells[2].paragraphs[0].runs[0]
                            run_firm.font.italic = True
                    
                    # Agregar resumen del progreso
                    doc.add_paragraph()
                    
                    # Calcular estadísticas del cronograma
                    total_fases = len(datos_fases)
                    fases_generadas = sum(1 for fase in datos_fases.values() if fase.get('generado'))
                    fases_firmadas = sum(1 for fase in datos_fases.values() if fase.get('firmado'))
                    
                    para_estadisticas = doc.add_paragraph()
                    para_estadisticas.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    stats_text = f"📊 Progreso: {fases_generadas}/{total_fases} generadas | {fases_firmadas}/{total_fases} firmadas"
                    run_stats = para_estadisticas.add_run(stats_text)
                    run_stats.font.bold = True
                    run_stats.font.size = Pt(12)
                    
                    # Determinar próxima fase
                    proxima_fase = None
                    for fase_key, fase_data in datos_fases.items():
                        if not fase_data.get('generado'):
                            proxima_fase = nombres_fases.get(fase_key, fase_key.title())
                            break
                    
                    if proxima_fase:
                        para_proxima = doc.add_paragraph()
                        para_proxima.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        proxima_text = f"🎯 Próxima fase: {proxima_fase}"
                        run_proxima = para_proxima.add_run(proxima_text)
                        run_proxima.font.bold = True
                        run_proxima.font.size = Pt(11)
                    
                else:
                    # No hay datos de fases
                    para_sin_fases = doc.add_paragraph("No hay información de fases disponible para este contrato")
                    para_sin_fases.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para_sin_fases.runs[0]
                    run.font.italic = True
            else:
                # Controlador de fases no disponible
                para_error = doc.add_paragraph("Controlador de fases no disponible")
                para_error.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para_error.runs[0]
                run.font.italic = True
                
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error insertando cronograma en Word: {e}")
            # Agregar mensaje de error en el documento
            para_error = doc.add_paragraph(f"Error generando cronograma: {e}")
            para_error.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para_error.runs[0]
            run.font.italic = True
    
    def _insertar_tabla_seguimiento_en_word(self, doc, datos_firmas: dict, firmantes_unicos: list):
        """Insertar tabla de seguimiento de firmas en el documento Word"""
        try:
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Título de la sección
            doc.add_page_break()
            doc.add_heading('📋 TABLA DE SEGUIMIENTO DE FIRMAS', level=1)
            
            doc.add_paragraph("Esta tabla muestra el estado de los documentos y firmas para cada fase del proyecto:")
            
            # Determinar fases según tipo de contrato (igual que en la tabla original)
            fases_mostrar = ['CREACION', 'INICIO', 'CARTASINVITACION', 'ADJUDICACION', 'CARTASADJUDICACION', 'CONTRATO', 
                           'REPLANTEO', 'ACTUACION', 'RECEPCION', 'FINALIZACION']
            
            # Crear tabla (Fase + Creación Doc + Fecha Última Firma + Firmantes)
            num_columnas = 3 + len(firmantes_unicos)
            tabla = doc.add_table(rows=1, cols=num_columnas)
            tabla.style = 'Table Grid'
            
            # Configurar encabezados
            headers = ['Fase', 'Creación Doc', 'Fecha Última Firma'] + [f'Firmante{i+1}' for i in range(len(firmantes_unicos))]
            
            for i, header in enumerate(headers):
                celda = tabla.cell(0, i)
                celda.text = header
                run = celda.paragraphs[0].runs[0]
                run.font.bold = True
                celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Llenar datos de cada fase
            for fase in fases_mostrar:
                row_cells = tabla.add_row().cells
                
                # Columna Fase
                row_cells[0].text = fase
                
                if fase in datos_firmas:
                    info_fase = datos_firmas[fase]
                    
                    # Columna Creación Doc
                    fecha_creacion = info_fase.get('fecha_creacion', '')
                    row_cells[1].text = fecha_creacion if fecha_creacion else "Sin datos"
                    
                    # Columna Fecha Última Firma
                    if info_fase.get('firmas'):
                        fechas_firmas = [f['fecha'] for f in info_fase['firmas'] if 'fecha' in f]
                        if fechas_firmas:
                            fecha_ultima = max(fechas_firmas)
                            row_cells[2].text = fecha_ultima
                        else:
                            row_cells[2].text = ""
                    else:
                        row_cells[2].text = ""
                    
                    # Columnas de firmantes
                    for col, firmante in enumerate(firmantes_unicos, 3):
                        if col < len(row_cells):
                            # Buscar si este firmante firmó esta fase
                            firma_encontrada = None
                            for firma in info_fase.get('firmas', []):
                                if firma['firmante'] == firmante:
                                    firma_encontrada = firma
                                    break
                            
                            if firma_encontrada:
                                dni_texto = f" - {firma_encontrada['dni']}" if 'dni' in firma_encontrada and firma_encontrada['dni'] else ""
                                row_cells[col].text = f"{firmante}{dni_texto}\n{firma_encontrada['fecha']}"
                            else:
                                row_cells[col].text = ""
                else:
                    # Fase sin datos
                    row_cells[1].text = "Sin datos"
                    row_cells[2].text = ""
                    for col in range(3, len(row_cells)):
                        row_cells[col].text = ""
            
            # Agregar leyenda
            doc.add_paragraph("\n📖 Leyenda:")
            doc.add_paragraph("• Creación Doc: Fecha cuando se generó el documento")
            doc.add_paragraph("• Fecha Última Firma: Fecha de la firma más reciente en esa fase")
            doc.add_paragraph("• Firmantes: Nombre, DNI y fecha de firma de cada firmante")
            
        except Exception as e:
            logger.error(f"[ControladorDocumentos] Error insertando tabla seguimiento en Word: {e}")
            # Agregar mensaje de error en el documento
            para_error = doc.add_paragraph(f"Error generando tabla de seguimiento: {e}")
            para_error.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para_error.runs[0]
            run.font.italic = True
    
    def _guardar_fichero_resumen(self, nombre_contrato: str, documento_word) -> str:
        """Guardar el documento Word de resumen en carpeta 10-otros de la obra"""
        from datetime import datetime
        
        try:
            # Crear nombre de archivo seguro para Word
            nombre_archivo = f"resumen_{nombre_contrato}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            nombre_archivo = "".join(c for c in nombre_archivo if c.isalnum() or c in '._-')
            
            # Crear directorio específico de la obra usando el nombre de carpeta correcto del JSON
            nombre_carpeta_real = self._obtener_nombre_carpeta_actual(nombre_contrato)
            directorio_obra = os.path.join(os.getcwd(), "obras", nombre_carpeta_real, "10-otros")
            
            # Crear el directorio si no existe
            os.makedirs(directorio_obra, exist_ok=True)
            logger.info(f"[ControladorDocumentos] Directorio creado/verificado: {directorio_obra}")
            
            # Ruta completa del archivo
            ruta_archivo = os.path.join(directorio_obra, nombre_archivo)
            
            # Guardar documento Word
            documento_word.save(ruta_archivo)
            logger.info(f"[ControladorDocumentos] Documento guardado en: {ruta_archivo}")
            
            return ruta_archivo
            
        except Exception as e:
            logger.warning(f"[ControladorDocumentos] Error guardando en carpeta obra, intentando ubicación de respaldo: {e}")
            
            # Si falla, usar directorio de reportes como respaldo
            directorio_reportes = os.path.join(os.getcwd(), "reportes")
            os.makedirs(directorio_reportes, exist_ok=True)
            
            ruta_archivo_respaldo = os.path.join(directorio_reportes, nombre_archivo)
            documento_word.save(ruta_archivo_respaldo)
            
            logger.info(f"[ControladorDocumentos] Documento guardado en ubicación de respaldo: {ruta_archivo_respaldo}")
            return ruta_archivo_respaldo